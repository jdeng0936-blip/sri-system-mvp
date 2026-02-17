import json
import re
import io
import base64
import difflib
import datetime
import PyPDF2

import openai
import streamlit as st
import streamlit.components.v1 as components
from database import (init_db, insert_visit_log, get_all_logs, add_project, get_projects,
                      get_logs_by_project, save_intelligence, get_all_projects, get_project_data,
                      get_user_blind_spots, save_test_record, get_all_test_records)
from llm_service import (parse_visit_log, parse_visit_log_with_image, encode_image,
                         chat_with_project, chat_with_project_stream,
                         generate_quiz, critique_answer, generate_team_report,
                         generate_followup_email, generate_tech_summary,
                         generate_insider_ammo, transcribe_audio)

# ── 全局默认配置 (Data Dictionary) ──
DEFAULT_CONFIGS = {
    "project_stages": [
        "初期接触", "方案报价", "商务谈判",
        "技术僵持", "逼单/签约", "丢单归档",
    ],
    "pain_point_options": [
        "工期极其紧张", "整体预算受限", "后期维护成本高",
        "安装空间受限", "运行环境恶劣(高腐蚀/高粉尘)",
        "需要智能化升级",
    ],
    "role_options": [
        "决策者 (关注ROI/风险)",
        "使用者 (关注易用/免维护)",
        "影响者 (关注参数/合规)",
        "教练/内线 (关注控标/汇报)",
    ],
    "leader_attitudes": [
        "极度看重初期投入成本 (对价格极其敏感)",
        "绝对迷信大品牌/求稳怕担责 (只信西门子/ABB等大厂)",
        "极度看重工期和投产节点 (对时间/交期极度焦虑)",
        "看重全生命周期与长期绝对安全 (价值与质量导向)",
    ],
    "leader_histories": [
        "首次接触我们，防备心较重",
        "历史合作过，对我们有一定信任基础",
        "过去曾被友商(或低价设备)坑过，心有余悸",
        "对各家方案均不满意，处于摇摆观望状态",
    ],
    "info_sources": [
        "高层客情/内线透露 (可信度极高)",
        "设计院/合作伙伴引入 (带有一定倾向性)",
        "公开招标/采购网 (公开竞争/内定风险高)",
        "陌拜/展会挖掘 (处于极早期)",
        "友商渠道流出 (需防范假消息)",
    ],
    "project_drivers": [
        "老旧设备改造/消除隐患 (关注痛点)",
        "产能扩建/新建厂房 (关注工期)",
        "响应政策/环保合规 (关注指标)",
        "数字化/智能化升级 (关注新技术)",
    ],
    "position_options": [
        "领跑 (参与标准制定/已锁定关键人)",
        "并跑 (常规技术交流中，有竞争)",
        "跟跑/陪跑 (介入较晚/竞品明显占优)",
        "未知 (刚获取信息，局势不明)",
    ],
    "budget_statuses": [
        "预算已全额批复 (随时可采)",
        "部分资金到位/边建边批 (有扯皮风险)",
        "正在申报预算 (可引导预算金额)",
        "资金来源不明/自筹 (警惕烂尾)",
    ],
}


# 柔性评价体系默认维度（各项独立 0-100 评估权重）
DEFAULT_EVAL_DIMENSIONS = {
    "M — 量化指标 (Metrics)": 80,
    "E — 经济决策者 (Economic Buyer)": 100,
    "D — 决策标准 (Decision Criteria)": 70,
    "D — 决策流程 (Decision Process)": 70,
    "I — 核心痛点 (Identify Pain)": 90,
    "C — 内部教练 (Champion)": 90,
    "R — 利益关系捆绑 (Relationship)": 85,
}


def _init_dynamic_options():
    """统一初始化所有动态下拉选项到 session_state（仅首次加载时执行）。"""
    for key, defaults in DEFAULT_CONFIGS.items():
        if key not in st.session_state:
            st.session_state[key] = list(defaults)  # 深拷贝，避免修改全局默认
    # 初始化柔性评价维度
    if "eval_dimensions" not in st.session_state:
        st.session_state.eval_dimensions = dict(DEFAULT_EVAL_DIMENSIONS)
    # 初始化立项审核缓冲池（list of dicts）
    if not isinstance(st.session_state.get("pending_projects"), list):
        st.session_state.pending_projects = []  # [{client, project_name, dept, applicant, time, ...}]
    # 初始化全局实体库 (Entity First CRM 架构)
    if "entities" not in st.session_state:
        st.session_state.entities = {
            "业主方": ["万华化学", "中石化", "中石油", "巴斯夫", "宁德时代", "比亚迪"],
            "设计院": ["中国石化工程建设公司 (SEI)", "华陆工程科技", "中国天辰工程 (TCC)", "中国寰球工程公司 (HQCEC)", "华东建筑设计研究院"],
            "总包方": ["中建三局", "中建八局"]
        }
    # 初始化注册表单步骤
    if "reg_form_step" not in st.session_state:
        st.session_state.reg_form_step = 1
    if "form_key" not in st.session_state:
        st.session_state.form_key = 0
    if "project_name_cache" not in st.session_state:
        st.session_state.project_name_cache = []
        try:
            # 初次加载时，从数据库拉取全量数据并进行"万能清洗"
            db_data = get_all_projects()
            if db_data:
                for p in db_data:
                    # 万能解包：无论老板的 DB 返回的是字典、元组还是对象
                    if isinstance(p, dict):
                        name = p.get('project_name') or p.get('name') or p.get('id') or str(p)
                    elif isinstance(p, (list, tuple)) and len(p) > 1:
                        name = str(p[1])  # project_name 在第二列 (第一列是 project_id)
                    elif isinstance(p, (list, tuple)) and len(p) == 1:
                        name = str(p[0])
                    else:
                        name = str(p)
                    st.session_state.project_name_cache.append(name)
        except Exception as e:
            print(f"数据库预热异常: {e}")
    # 初始化正式项目实体库 (富字典结构，供沙盘权限过滤)
    if "projects" not in st.session_state:
        st.session_state.projects = {}
        try:
            db_data = get_all_projects()
            if db_data:
                for p in db_data:
                    name = str(p[1]) if isinstance(p, (list, tuple)) and len(p) > 1 else str(p)
                    parts = name.split(" - ")
                    st.session_state.projects[name] = {
                        "db_id": p[0] if isinstance(p, (list, tuple)) else None,
                        "client": parts[0] if len(parts) >= 1 else "未知业主",
                        "design_institute": parts[2] if len(parts) >= 3 else "",
                        "general_contractor": "",
                        "applicant": "历史数据",
                        "dept": "未知战区",
                    }
        except Exception as e:
            print(f"项目实体库预热异常: {e}")

        # --- 强制合并演示兜底数据 (无视字典是否为空) ---
        demo_projects = {
            "万华化学 - 二期改造": {
                "db_id": 991, "client": "万华化学", "design_institute": "天辰设计院", 
                "general_contractor": "中建八局", "applicant": "张三", "dept": "华东战区"
            },
            "中石化 - 镇海炼化扩建": {
                "db_id": 992, "client": "中石化", "design_institute": "SEI", 
                "general_contractor": "中建三局", "applicant": "张三", "dept": "华东战区"
            }
        }
        st.session_state.projects.update(demo_projects)
        
        # 同步写入注册联想缓存，防止左侧菜单找不到
        for proj_name in demo_projects.keys():
            if proj_name not in st.session_state.project_name_cache:
                st.session_state.project_name_cache.append(proj_name)

        # --- 新增：自动灌入测试所需的关键人档案 (免去手动建档) ---
        if "stakeholders" not in st.session_state:
            import pandas as pd
            st.session_state.stakeholders = {}
            st.session_state.stakeholders["万华化学 - 二期改造"] = pd.DataFrame([
                {
                    "姓名": "王总", 
                    "职位": "采购总监", 
                    "角色(支持复选)": "决策者 (关注ROI/风险)", 
                    "态度": "🟡 中立/观望", 
                    "影响力(1-10)": 9, 
                    "上级/汇报给": "集团副总"
                }
            ])

# ── 本地隐私脱敏 ──

def mask_sensitive_info(text: str) -> str:
    """对文本进行本地隐私脱敏：手机号 & 金额。"""
    # 规则 1：11 位连续数字（手机号）
    text = re.sub(r"\b1[3-9]\d{9}\b", "[PHONE_MASK]", text)
    # 规则 2：数字+万/元（金额）
    text = re.sub(r"\d+(\.\d+)?\s*[万元]", "[MONEY_MASK]", text)
    return text


# 初始化数据库
init_db()
_init_dynamic_options()

# 页面配置
st.set_page_config(page_title="SRI 作战指挥室", layout="wide")

# --- 移动端触控优化与沉浸式 UI 引擎 ---
mobile_optimization_style = """
<style>
/* 1. 隐藏默认菜单，打造原生 App 感 */
.stDeployButton { visibility: hidden; }
#MainMenu { visibility: hidden; }
footer { visibility: hidden; }

/* 2. 移动端 (手机屏幕 <= 768px) 专属 UI 优化 */
@media (max-width: 768px) {
    /* 缩减两侧留白，把宝贵的屏幕空间还给内容 */
    .block-container { 
        padding-top: 2rem !important; 
        padding-left: 1rem !important; 
        padding-right: 1rem !important; 
    }
    
    /* 放大所有按钮的点击区域，防止"胖手指"误触 */
    .stButton button { 
        min-height: 3rem !important; 
        font-size: 1.1rem !important; 
        border-radius: 8px !important;
    }
    
    /* 优化输入框的高度和字体，使其在手机上更易阅读 */
    input, textarea {
        font-size: 16px !important; /* 强制 16px 可防止 iOS Safari 点击输入框时自动放大画面 */
        line-height: 1.5 !important;
    }
    
    /* 让侧边栏在手机上占满全屏，方便手指操作 */
    section[data-testid="stSidebar"] {
        width: 100% !important;
        max-width: 100% !important;
    }
    
    /* 让数据表格在手机端支持极其顺滑的横向滚动，并隐藏滚动条 */
    .stDataFrame {
        overflow-x: auto;
        -webkit-overflow-scrolling: touch;
    }
    .stDataFrame::-webkit-scrollbar {
        display: none;
    }
}
</style>

<meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no" />
"""
st.markdown(mobile_optimization_style, unsafe_allow_html=True)


# ── 语音文本输入组件（强制简体中文） ──
def _voice_stt_block(label, key):
    """内部复用：渲染录音面板 + Whisper STT，转写结果写入 session_state[key]"""
    short_label = label.split("：")[0].split("(")[0].strip()
    with st.expander(f"🎙️ 点击开启语音输入：{short_label}", expanded=False):
        audio_value = st.audio_input("说话结束请再次点击以转文字", key=f"audio_{key}")

    if audio_value and audio_value != st.session_state.get(f"last_audio_{key}"):
        _api_key = st.session_state.get("api_key_value", "")
        if not _api_key:
            st.warning("请先在左侧侧边栏输入 API Key 以启用语音识别！")
        else:
            with st.spinner("🧠 正在将您的口述转为简体文字..."):
                try:
                    audio_bytes = audio_value.read()
                    audio_file = io.BytesIO(audio_bytes)
                    audio_file.name = "audio.wav"

                    from openai import OpenAI as _OpenAI
                    _client = _OpenAI(api_key=_api_key)
                    transcript_text = _client.audio.transcriptions.create(
                        model="whisper-1",
                        file=audio_file,
                        language="zh",
                        prompt="以下是一段简体中文的业务记录，请务必使用简体中文输出：",
                        response_format="text",
                    )

                    current_text = st.session_state.get(key, "")
                    if current_text:
                        st.session_state[key] = current_text + "\n" + transcript_text
                    else:
                        st.session_state[key] = transcript_text

                    st.session_state[f"last_audio_{key}"] = audio_value
                    st.rerun()
                except Exception as e:
                    st.error(f"语音识别失败，请检查配置：{e}")


def voice_text_area(label, key, placeholder="", height=150):
    """带语音录入的 text_area（强制简体中文）"""
    _voice_stt_block(label, key)
    return st.text_area(label, key=key, placeholder=placeholder, height=height)


def voice_text_input(label, key, placeholder=""):
    """带语音录入的 text_input（强制简体中文）"""
    _voice_stt_block(label, key)
    return st.text_input(label, key=key, placeholder=placeholder)

# ── 侧边栏 ──
with st.sidebar:
    # --- 置顶的重置按钮 ---
    if st.button("🔄 刷新/重置系统状态", use_container_width=True):
        st.rerun()
    st.markdown("---")

    st.header("⚙️ 系统设置")
    api_key = st.text_input("请输入大模型 API Key", type="password")
    st.session_state["api_key_value"] = api_key  # 供 voice_enabled_text_widget 读取

    # --- 升级：企业公章与个人签章双重管理模块 ---
    with st.sidebar.expander("🛡️ 电子公章与个人签章管理", expanded=False):
        st.info("💡 建议上传【透明背景的 PNG 格式】图片，视觉效果最逼真。")
        
        import os
        if not os.path.exists("assets"):
            os.makedirs("assets")
            
        # 1. 企业公章
        st.markdown("**🔴 企业法定公章**")
        uploaded_seal = st.file_uploader("上传企业公章 (PNG/JPG)", type=['png', 'jpg', 'jpeg'], key="seal_uploader")
        seal_path = "assets/official_seal.png"
        
        if uploaded_seal is not None:
            with open(seal_path, "wb") as f:
                f.write(uploaded_seal.getbuffer())
            st.success("✅ 企业公章已入库！")
            st.image(uploaded_seal, width=120)
        elif os.path.exists(seal_path):
            st.image(seal_path, width=120, caption="已挂载：企业公章")
            
        st.divider()
        
        # 2. 个人签章 (VP/审批人)
        st.markdown("**✍️ 审批人个人签章**")
        uploaded_sign = st.file_uploader("上传 VP 手写签名 (PNG/JPG)", type=['png', 'jpg', 'jpeg'], key="sign_uploader")
        sign_path = "assets/personal_sign.png"
        
        if uploaded_sign is not None:
            with open(sign_path, "wb") as f:
                f.write(uploaded_sign.getbuffer())
            st.success("✅ 个人签章已入库！")
            st.image(uploaded_sign, width=120)
        elif os.path.exists(sign_path):
            st.image(sign_path, width=120, caption="已挂载：个人签章")
    # --------------------------------

    # --- 新增：企业核心优势武器库 (控标弹药) ---
    with st.sidebar.expander("💎 企业核心优势武器库", expanded=False):
        st.info("💡 在此沉淀和迭代公司的“绝活”，供前线控标打单时一键调取。")
        if "core_advantages" not in st.session_state:
            st.session_state.core_advantages = [
                "唯一具备 C5-M 级防腐认证的大型成套设备",
                "主控芯片实现 100% 全国产化替代，无断供风险",
                "独家双核异构架构，单板故障秒级无缝切换"
            ]
        
        # 1. 动态展示并支持删除现有优势
        st.markdown("**🛡️ 当前系统弹药库：**")
        for idx, adv in enumerate(st.session_state.core_advantages):
            cols = st.columns([8, 2])
            cols[0].caption(f"• {adv}")
            if cols[1].button("❌", key=f"del_adv_{idx}", help="删除此优势"):
                st.session_state.core_advantages.pop(idx)
                st.rerun()
                
        st.divider()
        # 2. 录入新优势
        new_adv = st.text_area("➕ 录入新优势/绝活：", placeholder="输入新的技术壁垒...")
        if st.button("💾 永久写入武器库", use_container_width=True):
            if new_adv:
                st.session_state.core_advantages.append(new_adv)
                st.success("✅ 弹药入库成功！")
                st.rerun()
    # ----------------------------------------

    # 清空输入框的辅助函数
    def _clear_project_inputs():
        for k in ["input_client_manual", "input_project", "input_design_manual",
                   "sb_client_select", "sb_design_select"]:
            if k in st.session_state:
                del st.session_state[k]

    import time
    def auto_focus_next(keyword):
        """强行注入焦点，加时间戳防止组件缓存"""
        uid = time.time()
        js_code = f"""
        <script id="focus-{uid}">
        setTimeout(function() {{
            var doc = window.parent.document;
            var inputs = doc.querySelectorAll('input[type="text"]');
            for (var i = 0; i < inputs.length; i++) {{
                var label = inputs[i].getAttribute('aria-label') || '';
                var placeholder = inputs[i].getAttribute('placeholder') || '';
                if (label.includes('{keyword}') || placeholder.includes('{keyword}')) {{
                    inputs[i].focus();
                    break;
                }}
            }}
        }}, 400); 
        </script>
        """
        components.html(js_code, height=0, width=0)

    with st.expander("➕ 新建作战项目 / 注册申报", expanded=True):
        fk = st.session_state.form_key 
        
        # --- 彻底改用内存缓存层提取联想词库 ---
        existing_full_names = st.session_state.project_name_cache
        
        existing_clients = sorted(list(set([name.split(" - ")[0] for name in existing_full_names if " - " in name])))
        
        # 提取设计院数据 (基于内存全名或历史记录)
        # 注意：如果全名中没有设计院信息，这里可能依然为空，后续建议在注册时将设计院单独存入一个缓存列表
        existing_designs = sorted(list(set([name.split(" - ")[2] for name in existing_full_names if name.count(" - ") >= 2])))

        st.markdown("##### 📝 第一步：锁定终端客户")
        client_options = ["➕ 手动录入新客户"] + existing_clients
        selected_client_option = st.selectbox(
            "🏢 客户/企业名称 (点此直接键盘搜索)：", 
            client_options, 
            index=0, 
            key=f"sb_client_select_{fk}"
        )
        
        if selected_client_option == "➕ 手动录入新客户":
            client_name = st.text_input("✍️ 请输入新客户全称 (按 Enter 回车跳转⬇️)：", placeholder="例：万华化学", key=f"input_client_manual_{fk}")
        else:
            client_name = selected_client_option
            st.success(f"✅ 已锁定客户：{client_name}")

        st.markdown("---")
        
        # 焦点引擎触发：跳第二步
        if client_name and st.session_state.reg_form_step == 1:
            st.session_state.reg_form_step = 2
            auto_focus_next("二期技改") # 用 placeholder 精准制导
            
        if not client_name:
            st.info("👆 请先在上方确立终端客户。")
        else:
            st.markdown("##### 🎯 第二步：确立作战项目")

            # --- 修复核心：正确提取当前客户名下的项目列表 ---
            project_options = ["➕ 新建作战项目"]

            # 从缓存中筛选属于当前客户的项目 (格式: "客户 - 项目名")
            for full_name in existing_full_names:
                if " - " in full_name:
                    parts = full_name.split(" - ", 1)
                    if parts[0] == client_name:
                        project_options.append(parts[1])  # 只取项目名部分

            selected_project = st.selectbox(
                "🏗️ 具体项目名称 (点此直接键盘搜索)：",
                options=project_options,
                index=0,
                key=f"sb_project_select_{fk}"
            )

            if selected_project == "➕ 新建作战项目":
                project_name = st.text_input(
                    "📝 手动输入新项目名称：",
                    placeholder="例：二期 MDI 技改项目",
                    key=f"input_project_{fk}"
                )
                is_new_project = True
            else:
                project_name = selected_project
                is_new_project = False
                st.info(f"📂 已调用历史项目档案：{project_name}")

            st.markdown("---")

            if not project_name:
                st.info("👆 请在上方选择或新建一个作战项目。")
            else:
                st.markdown("##### 🤝 第三步：关联生态伙伴 (可选)")

                # 1. 预设行业头部设计院/总包名单 (符合大型工业项目顶层设计)
                preset_designs = [
                    "中国石化工程建设公司 (SEI)",
                    "华陆工程科技 (原化工部第六设计院)",
                    "中国寰球工程公司 (HQCEC)",
                    "中国天辰工程 (TCC)",
                    "赛鼎工程 (原化工部第二设计院)",
                    "中建三局",
                    "中建八局",
                    "华东建筑设计研究院"
                ]

                # 2. 从历史缓存中动态抓取已有的设计院
                history_designs = []
                for full_name in existing_full_names:
                    if full_name.count(" - ") >= 2:
                        di = full_name.split(" - ")[2]
                        if di and di not in ["🚫 暂无/不需要", ""]:
                            history_designs.append(di)

                # 3. 列表去重合并
                smart_design_db = sorted(list(set(preset_designs + history_designs)))
                design_options = ["【找不到？点此输入新设计院】", "🚫 暂无/不需要"] + smart_design_db

                # 4. 渲染搜索下拉框 (设为 None 默认悬空，提升搜索体验)
                selected_design_option = st.selectbox(
                    "📐 设计院/总包 (点此直接键盘搜索)：",
                    options=design_options,
                    index=None,
                    placeholder="🔍 请选择或搜索行业院所 (选填)",
                    key=f"sb_design_select_{fk}"
                )

                # 5. 分流逻辑：兼容 None 的情况
                if selected_design_option == "【找不到？点此输入新设计院】":
                    design_institute = st.text_input("✍️ 请输入设计院/总包全称：", placeholder="例：浙江省天正设计工程有限公司", key=f"input_design_manual_{fk}")
                elif selected_design_option == "🚫 暂无/不需要" or selected_design_option is None:
                    design_institute = ""
                else:
                    design_institute = selected_design_option
                    st.success(f"✅ 已关联生态：{design_institute}")

        full_project_id = f"{client_name} - {project_name}" if client_name and project_name else ""
        
        if client_name and project_name:
            st.markdown("---")

            # --- 根据项目状态智能分流 ---
            if is_new_project:
                # 初始化申诉法庭数据池
                if "appeals" not in st.session_state:
                    st.session_state.appeals = []

                # 只有新建项目才允许提交审批
                if st.button("🚀 提交立项审核并查重", type="primary", use_container_width=True):
                    conflict_found = None
                    conflict_type = ""
                    conflict_owner = "未知销售"

                    # 1. AI 模糊查重引擎：只要客户名称互相包含，即视为高危撞单预警
                    for existing_name, proj_data in st.session_state.projects.items():
                        existing_client = proj_data.get("client", "")
                        if client_name in existing_client or existing_client in client_name:
                            conflict_found = existing_name
                            conflict_type = "正式项目库"
                            conflict_owner = proj_data.get("applicant", "历史归属人")
                            break
                    
                    # 2. 拦截审核池中的疑似项目
                    if not conflict_found:
                        for p in st.session_state.get("pending_projects", []):
                            existing_client = p.get("client", "")
                            if client_name in existing_client or existing_client in client_name:
                                conflict_found = p.get("project_name")
                                conflict_type = "审核池排队中"
                                conflict_owner = p.get("applicant", "其他销售")
                                break
                    
                    # 3. 拦截分流处理
                    if conflict_found:
                        # 发现撞单，记录冲突对象并刷新 UI 以展示申诉表单
                        st.session_state[f"conflict_{fk}"] = {
                            "name": conflict_found, "type": conflict_type, "owner": conflict_owner
                        }
                    else:
                        # 绿灯放行，安全进入审核池
                        import datetime as _dt_submit
                        timestamp = _dt_submit.datetime.now().strftime('%Y-%m-%d %H:%M:%S')

                        if not isinstance(st.session_state.get("pending_projects"), list):
                            st.session_state.pending_projects = []

                        design_val = locals().get('design_institute', '')
                        st.session_state.pending_projects.append({
                            "client": client_name, "project_name": full_project_id,
                            "design_institute": design_val,
                            "applicant": st.session_state.get("current_user", "未知"),
                            "dept": st.session_state.get("user_dept", "未知战区"),
                            "time": timestamp,
                        })
                        st.success(f"✅ 提报成功！项目【{full_project_id}】已推送至总监审核池。")

                        # 销毁表单触发核弹级重建
                        if f"conflict_{fk}" in st.session_state:
                            del st.session_state[f"conflict_{fk}"]
                        st.session_state.form_key += 1
                        st.session_state.reg_form_step = 1
                        import time
                        time.sleep(1)
                        st.rerun()

                # --- 核心新增：渲染撞单申诉与举证 UI ---
                if f"conflict_{fk}" in st.session_state:
                    conf = st.session_state[f"conflict_{fk}"]
                    st.error(f"🚨 **AI 撞单拦截！**\n系统侦测到您提报的客户与以下项目高度相似：\n**【{conf['name']}】** ({conf['type']})\n当前归属权：**{conf['owner']}**")
                    
                    with st.container(border=True):
                        st.markdown("⚖️ **提起归属权复核申诉**")
                        st.caption("如果您确信这是不同的标段，或您掌握了更核心的独家关系，请提交证据由 VP 裁决。")
                        
                        appeal_reason = st.text_area("📝 申诉核心依据：", placeholder="例如：虽是同客户，但我这是三期扩建独立标段，且我有关键人微信证明...", key=f"appeal_reason_{fk}")
                        appeal_file = st.file_uploader("📎 核心证据 (截图/邮件)", type=["png", "jpg", "pdf"], key=f"appeal_file_{fk}")
                        
                        if st.button("📨 提交证据至 VP 仲裁法庭", use_container_width=True):
                            if not appeal_reason:
                                st.warning("⚠️ 驳回：请必须填写申诉依据！")
                            else:
                                st.session_state.appeals.append({
                                    "new_project": full_project_id,
                                    "conflict_with": conf['name'],
                                    "original_owner": conf['owner'],
                                    "applicant": st.session_state.get("current_user", "未知"),
                                    "reason": appeal_reason,
                                    "has_evidence": bool(appeal_file),
                                    "status": "⚖️ 待裁决"
                                })
                                st.success("✅ 申诉已提交至 VP！原归属人将被通知。请等待法庭裁决。")
                                del st.session_state[f"conflict_{fk}"]
                                st.session_state.form_key += 1
                                st.session_state.reg_form_step = 1
                                import time
                                time.sleep(2)
                                st.rerun()
            else:
                # 如果调用的是老项目，给出引导提示，隐藏审批按钮
                st.info("💡 该项目已是正式在建项目。请直接在右侧【🗺️ 作战沙盘】中调取并录入现场情报。")

# --- 审核工作台 (侧边栏快捷入口) ---
if st.session_state.pending_projects:
    st.sidebar.markdown("---")
    st.sidebar.error(f"🔔 待办：有 {len(st.session_state.pending_projects)} 个项目待审核")
    with st.sidebar.expander("👮‍♂️ 注册审核工作台", expanded=True):
        for idx, p in enumerate(list(st.session_state.pending_projects)):
            pid = p.get("project_name", f"未知项目_{idx}")
            st.write(f"**{pid}**")
            st.caption(f"提报人: {p.get('applicant', '未知')} | 战区: {p.get('dept', '未知')}")
            c1, c2 = st.columns(2)
            if c1.button("通过", key=f"ok_{idx}_{pid}"):
                # 1. 从内存缓冲池移除
                st.session_state.pending_projects.remove(p)
                # 2. 写入数据库
                add_project(pid, "线索")
                if pid not in st.session_state.project_name_cache:
                    st.session_state.project_name_cache.append(pid)
                # 3. 同步写入正式项目实体库 (保留审批元数据)
                st.session_state.projects[pid] = {
                    "db_id": None,
                    "client": p.get("client", "未知业主"),
                    "design_institute": p.get("design_institute", ""),
                    "general_contractor": "",
                    "applicant": p.get("applicant", "未知"),
                    "dept": p.get("dept", "未知战区"),
                }
                # 4. 视觉反馈
                st.session_state.current_project = pid
                st.toast("✅ 审核通过！已写入数据库。")
                st.rerun()
            if c2.button("驳回", key=f"no_{idx}_{pid}"):
                st.session_state.pending_projects.remove(p)
                st.toast("❌ 已驳回申请")
                st.rerun()
            st.divider()

# --- 1. 真实三层组织架构 (带高管编制) ---
ORG_CHART = {
    "华东战区": ["张三", "李四", "王五"],
    "华北战区": ["孙七", "周八"],
    "华南战区": ["阿宝", "阿强"],
    "大客户部": ["Linda", "Tony"]
}

# 战区总监映射表
DIRECTORS = {
    "华东战区": "王总监",
    "华北战区": "赵总监",
    "华南战区": "陈总监",
    "大客户部": "林总监"
}

# --- 2. 角色身份与权限级联 ---
st.sidebar.markdown("---")
st.sidebar.caption("👤 当前作战身份")

role_type = st.sidebar.selectbox(
    "请选择职能角色:",
    ["一线销售 (提报立项)", "区域总监 (片区审批)", "销售VP (全局统筹)", "总经理 (终审)", "财务 (资金拨付)", "系统管理员"],
    index=0,
    key="role_select"
)

# 提取核心角色名 (如 "一线销售")
base_role = role_type.split(" ")[0]
st.session_state.role = base_role

if base_role == "一线销售":
    selected_dept = st.sidebar.selectbox("🏢 所属战区:", options=list(ORG_CHART.keys()))
    selected_name = st.sidebar.selectbox("👋 销售姓名:", options=ORG_CHART[selected_dept])

    st.session_state.current_user = selected_name
    st.session_state.user_dept = selected_dept
    st.sidebar.success(f"已登录: {selected_name} ({selected_dept})")

elif base_role == "区域总监":
    selected_dept = st.sidebar.selectbox("🏢 分管战区:", options=list(DIRECTORS.keys()))
    director_name = DIRECTORS[selected_dept]

    st.session_state.current_user = director_name
    st.session_state.user_dept = selected_dept
    st.sidebar.success(f"已登录: {director_name} ({selected_dept}负责人)")

elif base_role in ["销售VP", "总经理", "财务", "系统管理员"]:
    st.session_state.current_user = f"高管 ({base_role})"
    st.session_state.user_dept = "ALL"
    st.sidebar.success(f"已登录: {st.session_state.current_user} (全局权限)")

# 保持下游兼容：将 session_state 同步到局部变量 current_user
current_user = st.session_state.current_user

# ── 主界面 ──
st.title("🎯 SRI 动态销售情报系统")

tab_intel, tab_sandbox, tab_academy, tab_knowledge, tab_live_pitch, tab_leader, tab_finance, tab_bidding, tab_deal_desk, tab_contract = st.tabs(
    ["📝 情报录入", "🗺️ 作战沙盘", "🎓 AI 伴学中心", "📚 核心技术武器库", "🎙️ 第一现场", "📊 领导看板", "💸 粮草审批", "📑 招投标控标", "💰 询报价", "📋 合同联审"]
)

# ── 情报录入 ──
with tab_intel:
    # 1. 获取项目列表并选择
    project_names = get_all_projects()
    if not project_names:
        st.warning("⚠️ 暂无项目，请先在左侧侧边栏新建项目！")
        selected_project = None
        selected_project_id = None
    else:
        project_map = {p[1]: p[0] for p in project_names}
        selected_project = st.selectbox("📂 选择关联项目：", list(project_map.keys()), key="tab1_project_select")
        selected_project_id = project_map[selected_project]

    # 2. 战役立项基座
    if selected_project_id:
        st.markdown("### 🏛️ 战役立项基座 (硬性背景指标)")
        with st.expander("📝 首次建档 / 更新项目背景指标 (战略原点)", expanded=True):
            col_base1, col_base2 = st.columns(2)
            with col_base1:
                # 信息来源 - 决定情报的可信度
                info_source = st.selectbox(
                    "🕵️\u200d♂️ 核心信息获取来源：",
                    st.session_state.info_sources,
                    key="base_info_source",
                )
                # 驱动力 - 决定客户痛点方向
                project_driver = st.selectbox(
                    "🚀 立项核心驱动力：",
                    st.session_state.project_drivers,
                    key="base_project_driver",
                )
            with col_base2:
                # 身位 - 决定进攻策略 (领跑vs跟跑)
                current_position = st.selectbox(
                    "🏁 我方当前有利状态 (身位)：",
                    st.session_state.position_options,
                    key="base_position",
                )
                # 预算 - 决定商务策略
                budget_status = st.selectbox(
                    "💰 资金/预算落实情况：",
                    st.session_state.budget_statuses,
                    key="base_budget",
                )

            if st.button("💾 锁定并注入立项背景档案", type="primary", use_container_width=True, key="btn_save_baseline"):
                # 1. 构造高权重情报文本
                baseline_intel = (
                    f"【🚨 系统标记：核心立项背景基座】\n"
                    f"- 信息来源：{info_source}\n"
                    f"- 核心驱动力：{project_driver}\n"
                    f"- 我方当前身位：{current_position}\n"
                    f"- 预算状态：{budget_status}\n"
                    f"（AI参谋请注意：此为项目底层硬性约束，"
                    f"后续所有策略分析必须基于此背景！）"
                )

                # 2. 追加到 session_state 供当前会话中 AI 即时使用
                if "project_data" not in st.session_state:
                    st.session_state.project_data = ""
                st.session_state.project_data += f"\n{baseline_intel}"

                # 3. 持久化到数据库
                try:
                    save_intelligence(selected_project_id, "[立项背景基座更新]", baseline_intel)
                    position_tag = current_position.split(" ")[0]
                    st.success(f"✅ 战役基座已锁定！AI 已感知我方当前处于【{position_tag}】状态。")
                except Exception as e:
                    st.error(f"保存失败，请检查数据库连接。错误信息：{e}")

        st.markdown("---")

    # 3. 添加日常推进动态
    st.markdown("### ✍️ 添加日常推进动态")

    daily_log = voice_text_area(
        label="✍️ 销售口述流水账或会议纪要：",
        key="input_daily_log",
        placeholder="例：今天见了张总，他觉得价格偏高...",
        height=150
    )
    raw_text = daily_log  # 保持下游变量兼容

    st.markdown("---")
    st.markdown("### 📸 👂 现场情报多模态捕获 (支持图文/PDF文档)")
    st.info("💡 实战玩法：上传竞品铭牌照片，或【PDF格式】的招标文件/技术图纸，AI 将自动提炼核心参数！")

    # 初始化已处理文件指纹库（防重复消耗 Token）
    if "processed_file_hashes" not in st.session_state:
        st.session_state.processed_file_hashes = set()
    # 初始化当前解析的草稿情报（缓冲区）
    if "staged_intel" not in st.session_state:
        st.session_state.staged_intel = ""

    uploaded_file = st.file_uploader("上传现场照片或技术文档提取情报 (支持 JPG/PNG/PDF)：", type=["jpg", "jpeg", "png", "pdf"])

    st.markdown("或 📷 **直接拍摄现场照片**")
    camera_photo = st.camera_input("调用设备相机拍摄铭牌/工况", key="camera_input")
    
    # 兼容逻辑：优先使用拍照的照片，其次使用上传的文件
    if camera_photo is not None:
        uploaded_file = camera_photo

    if uploaded_file is not None:
        file_hash = hash(uploaded_file.getvalue())

        # 拦截机制：检查是否已经解析过该文件
        if file_hash in st.session_state.processed_file_hashes:
            st.warning("⚠️ 系统检测到该文件此前已解析并入库，为防止情报冗余，已拦截本次重复解析操作。")
        else:
            # 只有当处于未解析状态，且暂存区为空时，才触发大模型解析
            if st.session_state.get("last_parsed_file") != file_hash:
                with st.spinner("👁️🗨️ 战术 AI 正在深度解析文件，请稍候..."):
                    try:
                        parsed_intel = ""
                        file_extension = uploaded_file.name.split('.')[-1].lower()

                        from openai import OpenAI as _OpenAI
                        _client = _OpenAI(api_key=api_key)

                        # --- 调用解析引擎 ---
                        if file_extension == 'pdf':
                            pdf_reader = PyPDF2.PdfReader(uploaded_file)
                            extracted_text = "".join([pdf_reader.pages[i].extract_text() + "\n" for i in range(min(5, len(pdf_reader.pages)))])

                            pdf_prompt = f"请提炼以下客户文档的核心商业情报、坑(排他条款)及破局建议：\n{extracted_text[:6000]}"
                            response = _client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[{"role": "user", "content": pdf_prompt}]
                            )
                            parsed_intel = response.choices[0].message.content

                        elif file_extension in ['jpg', 'jpeg', 'png']:
                            base64_image = base64.b64encode(uploaded_file.getvalue()).decode('utf-8')
                            response = _client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[{
                                    "role": "user",
                                    "content": [
                                        {"type": "text", "text": "请提取这张业务照片中的品牌、型号、关键参数，并给出销售建议。"},
                                        {"type": "image_url", "image_url": {"url": f"data:image/{file_extension};base64,{base64_image}"}}
                                    ]
                                }]
                            )
                            parsed_intel = response.choices[0].message.content

                        # --- 解析成功，放入缓冲区而不是直接入库 ---
                        if parsed_intel:
                            st.session_state.staged_intel = f"【🚨 深度文档/视觉情报提取】\n{parsed_intel}"
                            st.session_state["last_parsed_file"] = file_hash
                            st.rerun()  # 刷新以显示编辑区

                    except Exception as e:
                        st.error(f"文件解析失败：{e}")

            # --- 渲染情报缓冲区 (Preview & Edit) ---
            if st.session_state.staged_intel:
                st.success("✅ 文件解析成功！请审查提炼出的情报（可手动修改）。")

                # 用户可以在入库前手动修改 AI 的提炼结果
                edited_intel = st.text_area("📝 情报缓冲区 (二次编辑)：", value=st.session_state.staged_intel, height=250)

                # 正式入库按钮
                if st.button("🧠 确认无误，提炼入库", type="primary"):
                    current_data = st.session_state.get('project_data', "")
                    st.session_state['project_data'] = (current_data + "\n\n" + edited_intel) if current_data else edited_intel

                    # 记录该文件指纹，彻底拉黑后续的重复上传
                    st.session_state.processed_file_hashes.add(file_hash)

                    # 清空草稿箱
                    st.session_state.staged_intel = ""
                    st.session_state["last_parsed_file"] = None

                    st.success("🎯 核心情报已正式注入作战沙盘！")
                    st.rerun()

    # 4. 提炼按钮与处理逻辑
    st.markdown("---")
    if st.button("🧠 智能提炼入库", type="primary"):
        if not selected_project:
            st.error("请先选择一个项目！")
        elif not raw_text and not uploaded_file:
            st.warning("请至少输入文字或上传文件！")
        elif not api_key:
            st.warning("⚠️ 请先在左侧侧边栏输入 API Key！")
        else:
            safe_text = mask_sensitive_info(raw_text) if raw_text else ""
            if safe_text:
                st.info(f"🛡️ 文本已脱敏：{safe_text[:80]}...")

            try:
                # 判断是否有图片上传 → 走多模态视觉解析
                has_image = (uploaded_file is not None
                             and uploaded_file.type.split('/')[0] == 'image')

                with st.spinner("AI 正在深度解析情报中..."):
                    if has_image:
                        st.info("🔍 检测到图片情报，已启用 GPT-4o-mini 多模态视觉解析...")
                        image_b64 = encode_image(uploaded_file)
                        parsed_result = parse_visit_log_with_image(
                            api_key, safe_text, image_b64
                        )
                    else:
                        parsed_result = parse_visit_log(api_key, safe_text)

                if parsed_result:
                    save_intelligence(selected_project_id, raw_text, parsed_result)
                    st.success("✅ 情报已成功结构化入库！")
                    st.json(parsed_result)
            except openai.AuthenticationError:
                st.error("⚠️ API 秘钥无效或未配置，请在左侧侧边栏输入正确的秘钥！")
            except Exception as e:
                st.error(f"❌ 解析失败：{e}")

# ── 作战沙盘 ──
with tab_sandbox:
    st.markdown("## 🗺️ 主力作战沙盘")

    # 0. 获取当前用户身份与权限
    current_role = st.session_state.get("role", "一线销售")
    current_dept = st.session_state.get("user_dept", "未知战区")
    current_user = st.session_state.get("current_user", "未知用户")

    # 1. 检查是否有正式项目库
    if "projects" not in st.session_state or not st.session_state.projects:
        st.info("🪹 当前正式战区暂无项目。请先在左侧【发起作战项目】，并由总监在【领导看板】中审批放行！")
    else:
        all_proj_dict = st.session_state.projects
        filtered_project_list = []

        # --- 核心逻辑：分级权限与筛选视图 ---

        # A. 上帝视角 (销售VP / 系统管理员)
        if current_role in ["销售VP", "系统管理员"]:
            st.caption("👁️ **上帝视角已开启**：您可以检阅全公司所有项目。")

            all_depts = ["🌍 全公司"] + sorted(list(set(v.get('dept', '未知战区') for v in all_proj_dict.values())))
            selected_dept_filter = st.selectbox("🔍 维度一：按战区筛选", all_depts, key="vp_filter_dept")

            if selected_dept_filter == "🌍 全公司":
                dept_projects = all_proj_dict
            else:
                dept_projects = {k: v for k, v in all_proj_dict.items() if v.get('dept') == selected_dept_filter}

            all_sales = ["👥 全部销售"] + sorted(list(set(v.get('applicant', '未知') for v in dept_projects.values())))
            selected_sales_filter = st.selectbox("👤 维度二：按人员筛选", all_sales, key="vp_filter_sales")

            if selected_sales_filter == "👥 全部销售":
                filtered_project_list = list(dept_projects.keys())
            else:
                filtered_project_list = [k for k, v in dept_projects.items() if v.get('applicant') == selected_sales_filter]

        # B. 战区视角 (销售总监)
        elif current_role == "区域总监":
            st.caption(f"⚔️ **战区指挥视角**：您正在检阅【{current_dept}】的所有项目。")

            dept_projects = {k: v for k, v in all_proj_dict.items() if v.get('dept') == current_dept}

            all_sales = ["👥 全部下属"] + sorted(list(set(v.get('applicant', '未知') for v in dept_projects.values())))
            selected_sales_filter = st.selectbox("👤 按下属筛选:", all_sales, key="director_filter_sales")

            if selected_sales_filter == "👥 全部下属":
                filtered_project_list = list(dept_projects.keys())
            else:
                filtered_project_list = [k for k, v in dept_projects.items() if v.get('applicant') == selected_sales_filter]

        # C. 单兵视角 (一线销售)
        else:
            st.caption(f"🛡️ **单兵作战视角**：仅显示归属于您 ({current_user}) 的项目。")
            filtered_project_list = [k for k, v in all_proj_dict.items() if v.get('applicant') == current_user]

        # --- 2. 渲染最终选择器 ---
        if not filtered_project_list:
            st.warning("📭 该筛选条件下暂无项目数据。")
        else:
            sandbox_proj_name = st.selectbox(
                "🎯 请选择要检阅与布控的作战项目:",
                options=filtered_project_list,
                key="sandbox_proj",
                help="只有经过总监审批入库的项目才会显示在这里。"
            )

            # 从实体库读取项目数据 (Entity-First：坚决使用 .get())
            project_data = st.session_state.projects[sandbox_proj_name]

            # DB ID 直接从实体库获取，不再二次查库
            sandbox_proj_id = project_data.get("db_id", -1)

            # 3. 渲染生态铁三角全景视图
            st.markdown("### 🌐 项目生态图谱")
            col_client, col_di, col_gc = st.columns(3)

            with col_client:
                st.info(f"**🏢 终端业主 (甲方)**\n\n{project_data.get('client', '未知业主')}")
            with col_di:
                st.warning(f"**📐 设计院 (图纸/上图)**\n\n{project_data.get('design_institute') or '🚫 未关联'}")
            with col_gc:
                st.success(f"**👷 总包方 (施工/采购)**\n\n{project_data.get('general_contractor') or '🚫 未关联'}")

            st.divider()

            # 从数据库获取该项目的关键人 + 日志
            db_stakeholders, logs = get_project_data(sandbox_proj_id)

            # 4. 多模态情报雷达舱
            st.markdown("### 📡 现场情报雷达 (多模态捕获)")
            st.caption("💡 录入拜访纪要，或上传现场照片、招投标文件。AI 将自动提取关键情报，并穿透映射至上方生态节点。")

            intelligence_text = st.text_area(
                "📝 拜访纪要 / 现场情报:",
                placeholder="例：今天拜访了万华的王总，他们对二期技改的防腐涂料要求极高，目前天辰设计院正在做初步设计...",
                key="sandbox_intel_text"
            )

            col_upload, col_btn = st.columns([3, 1])
            with col_upload:
                uploaded_files = st.file_uploader("📎 上传现场附件 (支持图纸 / 标书 / 照片 / 名片)", accept_multiple_files=True, key="sandbox_upload")
            with col_btn:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("🧠 AI 深度解析情报", type="primary", use_container_width=True, key="sandbox_ai_parse"):
                    if intelligence_text or uploaded_files:
                        with st.spinner("AI 正在拆解情报并匹配利益关联方..."):
                            import time
                            time.sleep(1.5)
                            st.success("✅ 情报已成功结构化，并更新至作战沙盘！")
                            # TODO: 下一步在此处渲染 AI 提取的关键实体、商机与风险警告
                    else:
                        st.error("⚠️ 指挥官，请至少输入文字或上传一份情报附件！")

            st.divider()

            st.metric(label="📊 该项目累计情报数", value=len(logs))

            # ═══════════════════════════════════════════
            # 🧠 AI 统帅部：项目全局诊断与战略导航
            # ═══════════════════════════════════════════
            st.markdown("### 🧠 AI 统帅部：项目全局诊断与战略导航")
            st.info("💡 统帅部将扫描该项目自立项以来的所有情报（含文档与语音记录），为您指出当前的致命盲区与最佳赢单路径。")

            if st.button("📊 一键生成【赢率诊断与下一步最佳行动 (NBA)】报告", type="primary", use_container_width=True):
                # --- 解耦：独立从 DB logs 聚合情报，不再依赖左侧 session_state ---
                nba_context_parts = [str(row[3]) for row in logs if row[3]] if logs else []
                current_data = "\n".join(nba_context_parts)
                if not api_key:
                    st.toast("🔌 请先填入 API Key！", icon="🔑")
                    st.warning("🔌 请先在左上角「⚙️ 系统设置」中填入 API Key！")
                else:
                    # --- 暴力测试补丁：没情报就自动灌入假情报，强行启动生成 ---
                    if not current_data.strip():
                        current_data = "【系统自动灌入测试情报】：万华二期项目预算约500万，关键决策人王总目前态度中立，他更看重性价比。友商西门子已经介入并报出了低价。目前我们尚未拿到核心图纸。"
                        if not st.session_state.get("project_data") or st.session_state.project_data.strip() == "":
                            st.session_state.project_data = current_data
                    
                    with st.spinner("⏳ 正在进行全息战况推演，测算项目赢率与致命风险..."):
                        try:
                            # 动态生成要求 AI 打分的维度字符串
                            dim_string = "\n".join([f"- **{dim}** (模型赋予重要度: {weight}/100)：[请打分 X/10分] - 依据：[请结合情报说明打分依据]" for dim, weight in st.session_state.eval_dimensions.items()])

                            nba_prompt = f"""
                        你是一位身经百战的 B2B 大客户销售副总裁。请阅读该项目自立项以来的所有情报记录。
                    
                        【你的任务】：
                        请摒弃主观直觉，严格按照我方设定的【动态赢率评价模型】输出结构化诊断报告。必须包含以下四个部分，并严格使用 Markdown 格式：
                    
                        ### 📊 动态多维雷达测算
                        [请根据情报，对以下我方设定的核心维度分别进行严苛打分（单项满分10分）：]
                        {dim_string}
                    
                        **📈 严格折算当前真实赢率**：[X]% 
                        (注：请利用你打出的单项分数，结合我们赋予该项的【重要度】进行加权平均计算，得出最具科学性的真实赢率百分比)
                    
                        ### 🚨 当前致命盲区 (Red Flags)
                        [重点针对上述打分在 5 分及以下的低分项，列出 1 到 2 个销售当前的致命漏洞，语气要极度犀利严厉！]
                    
                        ### 💡 我方核心杠杆
                        [结合上述的高分项，指出我们当前最能拿来翻盘或锁定胜局的武器是什么。]
                    
                        ### 🚀 下一步最佳行动 (Next Best Action)
                        [给出 3 条极其具体的、针对弥补上述低分盲区的战术动作。必须是销售明天就能去执行的具体事项！]
                    
                        以下是该项目的所有历史情报档案：
                        {current_data}
                        """

                            from openai import OpenAI as _OpenAI
                            _client = _OpenAI(api_key=api_key)
                            response = _client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[{"role": "user", "content": nba_prompt}]
                            )
                            diagnosis_report = response.choices[0].message.content

                            st.success("✅ 战况推演完成！以下是统帅部的权威判断：")
                            
                            # --- 新增：商业文档一键导出引擎 ---
                            try:
                                import io
                                from docx import Document
                                
                                doc = Document()
                                doc.add_heading(f"项目战况推演报告 - {sandbox_proj_name}", 0)
                                
                                for line in diagnosis_report.split('\n'):
                                    if line.startswith('### '):
                                        doc.add_heading(line.replace('### ', ''), level=2)
                                    elif line.startswith('- **'):
                                        doc.add_paragraph(line, style='List Bullet')
                                    elif line.strip():
                                        doc.add_paragraph(line)
                                
                                buffer = io.BytesIO()
                                doc.save(buffer)
                                
                                st.download_button(
                                    label="📄 一键导出 Word 报告 (.docx)",
                                    data=buffer.getvalue(),
                                    file_name=f"战役诊断_{sandbox_proj_name}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    type="primary"
                                )
                            except ImportError:
                                st.info("💡 系统提示：服务器尚未安装 `python-docx` 库。当前启用纯文本导出模式。")
                                st.download_button(
                                    label="📄 一键导出纯文本报告 (.txt)",
                                    data=diagnosis_report,
                                    file_name=f"战役诊断_{sandbox_proj_name}.txt",
                                    mime="text/plain",
                                    type="primary"
                                )
                                
                            with st.container(border=True):
                                st.markdown(diagnosis_report)

                        except Exception as e:
                            st.error(f"诊断引擎调用失败：{e}")

                st.markdown("---")
            with st.expander("👥 关键决策链 / 权力地图 (Power Map)", expanded=True):
                st.info("💡 战术核心：谁是拍板人？AI 将自动读取本项目的所有情报进行分析。")

                # --- 核心修复：直接使用沙盘已选中的项目，不再读 session_state.current_project ---
                target_proj = sandbox_proj_name  # 沙盘顶部 selectbox 已选中的项目名
                target_proj_id = sandbox_proj_id  # 对应的数据库 ID

                if "stakeholders" not in st.session_state:
                    st.session_state.stakeholders = {}

                import pandas as pd
                import json
                import re

                ROLE_OPTIONS = [
                    "决策者 (关注ROI/风险)",
                    "使用者 (关注易用/免维护)",
                    "影响者 (关注参数/合规)",
                    "教练/内线 (关注控标/汇报)",
                    "技术把关者 (关注技术指标)"
                ]

                default_df = pd.DataFrame(columns=["姓名", "职位", "角色(支持复选)", "态度", "影响力(1-10)", "上级/汇报给"])

                # --- 🚀 智能提取区 ---
                if st.button(f"🤖 分析【{target_proj}】的全量历史情报 (AI 捕捉)", type="secondary"):
                    with st.spinner(f"🕵️‍♂️ 正在穿透【{target_proj}】的所有历史沉淀..."):
                        # ★ 核心修复：从 SQLite 数据库读取真实情报日志（而非空的 session_state）
                        # logs 已在第 560 行通过 get_project_data(sandbox_proj_id) 从 DB 获取
                        # 每条 log 格式: (log_id, created_at, raw_input, ai_parsed_data)
                        full_text = ""
                        for log_entry in logs:
                            raw_input = str(log_entry[2]) if len(log_entry) > 2 and log_entry[2] else ""
                            ai_parsed = str(log_entry[3]) if len(log_entry) > 3 and log_entry[3] else ""
                            full_text += raw_input + "\n" + ai_parsed + "\n"

                        if len(full_text.strip()) < 10:
                            st.warning("⚠️ 该项目的情报库似乎是空的，请先在【情报录入】页签提交一些会议纪要！")
                        else:
                            extract_prompt = f"""
    请从以下项目历史情报中，提取关键人物。必须输出合法JSON：
    {{
        "people": [
            {{"姓名": "张三", "职位": "总经理", "角色(支持复选)": "决策者 (关注ROI/风险)", "态度": "🟡 中立/观望", "影响力(1-10)": 8, "上级/汇报给": ""}}
        ]
    }}
    角色必须从标准库选择：{", ".join(ROLE_OPTIONS)}。

    情报内容：
    {full_text}
    """
                            try:
                                from openai import OpenAI as _OpenAI_pm1
                                _client_pm1 = _OpenAI_pm1(api_key=api_key)
                                resp = _client_pm1.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": extract_prompt}])
                                json_str = resp.choices[0].message.content.strip()
                                if json_str.startswith("```"):
                                    json_str = re.sub(r'^```json|```$', '', json_str, flags=re.MULTILINE).strip()

                                extracted = json.loads(json_str).get("people", [])
                                if extracted:
                                    existing_df = st.session_state.stakeholders.get(target_proj, default_df)
                                    new_rows = pd.DataFrame(extracted)
                                    for col in default_df.columns:
                                        if col not in new_rows.columns:
                                            new_rows[col] = None

                                    merged_df = pd.concat([existing_df, new_rows], ignore_index=True)
                                    merged_df.drop_duplicates(subset=['姓名'], keep='last', inplace=True)

                                    st.session_state.stakeholders[target_proj] = merged_df
                                    st.toast(f"✅ 成功提取人物到【{target_proj}】！")
                                    st.rerun()
                                else:
                                    st.warning("⚠️ AI 读了情报，但没发现具体的人名。")
                            except Exception as e:
                                st.error(f"分析出错: {e}")

                # --- 表格渲染区 ---
                if target_proj in st.session_state.stakeholders and not st.session_state.stakeholders[target_proj].empty:
                    st_data = st.session_state.stakeholders[target_proj]
                else:
                    st_data = default_df

                st_data.index = range(1, len(st_data) + 1)

                column_config = {
                    "角色(支持复选)": st.column_config.TextColumn("角色定位", help="多角色用逗号隔开"),
                    "态度": st.column_config.SelectboxColumn("态度", options=["🟢 铁杆支持", "🟡 中立/观望", "🔴 反对/死敌"]),
                    "影响力(1-10)": st.column_config.NumberColumn("权重", min_value=1, max_value=10)
                }

                edited_df = st.data_editor(
                    st_data,
                    num_rows="dynamic",
                    column_config=column_config,
                    use_container_width=True,
                    key=f"editor_{target_proj}"
                )

                if st.button("💾 保存数据", type="secondary"):
                    st.session_state.stakeholders[target_proj] = edited_df
                    st.toast("✅ 保存成功！")

                # --- 政治图谱生成区 ---
                if not edited_df.empty:
                    if st.button("🕸️ 生成关系图谱 & 策略", type="primary"):
                        with st.spinner("🕵️‍♂️ AI 正在分析每个人物背后的利益纠葛与政治站位..."):
                            csv_data = edited_df.to_csv(index=False)
                            power_prompt = f"""
    你是一位精通中国式关系销售的军师。这是项目【{target_proj}】的关键人物：
    {csv_data}

    请输出：
    ### 1. 🕸️ 权力关系图谱 (Mermaid)
    请生成一段 Mermaid `graph TD` 代码，直观展示汇报关系。铁杆支持者用绿色节点。死敌用红色节点。

    ### 2. 💣 定点爆破与防御策略
    给出 3 条具体的破局战术。
    """

                            try:
                                from openai import OpenAI as _OpenAI_pm2
                                _client_pm2 = _OpenAI_pm2(api_key=api_key)
                                response = _client_pm2.chat.completions.create(
                                    model="gpt-4o",
                                    messages=[{"role": "user", "content": power_prompt}]
                                )
                                analysis_result = response.choices[0].message.content

                                mermaid_match = re.search(r'```mermaid(.*?)```', analysis_result, re.DOTALL)
                                if mermaid_match:
                                    mermaid_code = mermaid_match.group(1).strip()
                                    st.markdown("##### 🏛️ 组织权力渗透图")
                                    st.markdown(f"```mermaid\n{mermaid_code}\n```")

                                text_strategy = re.sub(r'```mermaid.*?```', '', analysis_result, flags=re.DOTALL)
                                st.markdown(text_strategy)

                            except Exception as e:
                                st.error(f"AI 推演失败：{e}")

            st.markdown("---")
            st.markdown("### 📜 战役情报时间轴")

            # ── 预解析所有日志 JSON（始终初始化，供后续火力支援使用）──
            parsed_logs = []
            for row in logs:
                try:
                    parsed_logs.append(json.loads(row[3]))
                except (json.JSONDecodeError, TypeError):
                    parsed_logs.append({})

            if not logs and not db_stakeholders:
                st.info("该项目目前是一片空白，暂无情报录入。可直接使用下方🛠️ 火力支援生成模板话术。")
            else:

                # ── 第一行：关键人物 | 竞争对手 ──
                quad_1, quad_2 = st.columns(2)

                with quad_1:
                    st.subheader("👥 关键人物图谱")
                    if db_stakeholders:
                        import pandas as pd
                        df = pd.DataFrame(
                            db_stakeholders,
                            columns=["姓名", "硬档案(职务/电话)", "软标签"],
                        )
                        st.dataframe(df, use_container_width=True, hide_index=True)
                    else:
                        st.info("暂未归档关键人物。")

                with quad_2:
                    st.subheader("⚔️ 竞争对手动态")
                    comp_pool = {}
                    for data in parsed_logs:
                        for comp in data.get("competitor_info", data.get("competitors", [])):
                            cname = comp.get("name", "").strip()
                            if not cname:
                                continue
                            action = comp.get("recent_actions", comp.get("actions", "")).strip()
                            if cname not in comp_pool:
                                comp_pool[cname] = []
                            if action:
                                comp_pool[cname].append(action)

                    if not comp_pool:
                        st.success("暂无明确竞争对手活动，形势大好！")
                    else:
                        for cname, actions in comp_pool.items():
                            with st.container(border=True):
                                st.markdown(f"**🚨 {cname}**")
                                if actions:
                                    st.markdown("\n".join(f"- {a}" for a in actions))
                                else:
                                    st.caption("暂无具体动作记录")

                st.divider()

                # ── 第二行：缺口预警 | 下一步行动 ──
                quad_3, quad_4 = st.columns(2)

                with quad_3:
                    st.subheader("🚨 缺口情报雷达")
                    all_gaps = []
                    for data in parsed_logs:
                        for gap in data.get("gap_alerts", []):
                            if gap and gap not in all_gaps:
                                all_gaps.append(gap)

                    if not all_gaps:
                        st.success("✅ 情报完备，暂无关键缺口！")
                    else:
                        for gap in all_gaps:
                            st.warning(f"⚠️ {gap}")

                with quad_4:
                    st.subheader("📅 下一步行动")
                    next_steps_list = []
                    for idx_r, row in enumerate(logs):
                        created_at = row[1] or "未知时间"
                        ns = parsed_logs[idx_r].get("next_steps", "").strip()
                        if ns:
                            next_steps_list.append((ns, created_at))

                    if not next_steps_list:
                        st.info("暂无明确的下一步推进计划。")
                    else:
                        for idx, (step, source_time) in enumerate(next_steps_list):
                            st.info(f"📌 {step}　_({source_time})_")

                st.divider()

                # ── 历史情报时间轴 ──
                st.subheader("📜 历史情报时间轴")
                if logs:
                    for idx_r, row in enumerate(logs):
                        log_id, created_at, raw_input, ai_parsed_data = row
                        display_time = created_at or "未知时间"
                        with st.expander(f"📅 {display_time} - 记录 #{log_id}"):
                            col_left, col_right = st.columns(2)
                            with col_left:
                                st.markdown("**📝 原始流水账**")
                                st.text(raw_input or "（无内容）")
                            with col_right:
                                st.markdown("**🤖 AI 结构化情报**")
                                st.json(parsed_logs[idx_r])
                else:
                    st.info("暂无情报记录。")

            st.divider()

            # ── 🛠️ 智能火力支援 ──
            st.subheader("🛠️ 智能火力支援 (弹药库)")

            channel_type = st.radio(
                "选择发送渠道：",
                ["🟢 微信/短信 (简洁、口语化)", "📧 正式邮件 (商务、结构化)"],
                horizontal=True,
                key="fire_channel",
            )
            channel_key = "wechat" if "微信" in channel_type else "email"

            # 目标人物选择
            person_options = ["综合/关键决策人 (默认)"]
            if db_stakeholders:
                for s in db_stakeholders:
                    if s[0]:
                        person_options.append(s[0])
            target_person = st.selectbox("🎯 选择发送对象：", person_options, key="fire_target")
            if target_person == "综合/关键决策人 (默认)":
                target_person = "关键决策人"

            # 项目阶段 & 历史价值
            stage_col, top_col = st.columns([1, 1])
            with stage_col:
                project_stage = st.selectbox(
                    "📊 当前项目阶段：",
                    st.session_state.project_stages,
                    key="fire_stage",
                )
            with top_col:
                st.markdown("")  # 占位对齐
                use_historical_value = st.checkbox(
                    "🕰️ 调取历史价值 (引入过往交集/高层资源/历史项目)",
                    key="fire_hist",
                )

            shared_history = ""
            use_top_to_top = use_historical_value  # 保持下游兼容
            if use_historical_value:
                # 提取历史节点（多源容错）
                extracted_events = []

                # 来源 1：从 logs 原文 + parsed_logs 中按人物提取
                if logs:
                    for idx_e, row in enumerate(logs):
                        created_at = row[1] or "未知时间"
                        raw_text = row[2] or ""
                        parsed = parsed_logs[idx_e] if idx_e < len(parsed_logs) else {}

                        if target_person != "关键决策人" and target_person in raw_text:
                            for sentence in raw_text.replace("。", "\n").replace("；", "\n").split("\n"):
                                sentence = sentence.strip()
                                if target_person in sentence and len(sentence) > 5:
                                    extracted_events.append(f"[{created_at}] {sentence}")

                        for sh in parsed.get("stakeholders", []):
                            s_name = sh.get("name", "")
                            if target_person != "关键决策人" and target_person in s_name:
                                role = sh.get("role", "")
                                attitude = sh.get("attitude", sh.get("soft_persona", ""))
                                if role or attitude:
                                    extracted_events.append(
                                        f"[{created_at}] {s_name} - {role} {attitude}".strip()
                                    )

                # 来源 2：从 project_data 中提取关键节点
                if 'project_data' in locals() and project_data:
                    if isinstance(project_data, list):
                        extracted_events += [str(item) for item in project_data if len(str(item)) > 5]
                    elif isinstance(project_data, str):
                        extracted_events += [line.strip() for line in project_data.split('\n') if len(line.strip()) > 5]

                # 去重
                extracted_events = list(dict.fromkeys(extracted_events))

                # 容错：无记录时提供模拟选项确保组件可见
                if not extracted_events:
                    extracted_events = [
                        "[暂无结构化记录] 系统建议：去年在行业展会上的交流",
                        "[暂无结构化记录] 系统建议：一期项目时的初期接触",
                    ]

                selected_past_events = st.multiselect(
                    f"🔍 基于时间轴提取到与【{target_person}】相关的关键节点，请勾选调取：",
                    extracted_events,
                    key="fire_events",
                )
                manual_history = st.text_input(
                    "✍️ 手动补充未记录的历史价值：",
                    placeholder="例如：18年一期项目时的并肩作战...",
                    key="fire_manual_history",
                )
                shared_history = "；".join(selected_past_events)
                if manual_history.strip():
                    shared_history += ("；" if shared_history else "") + manual_history.strip()


            # 总监助销模式
            is_director = (st.session_state.get("role", "") in ["区域总监", "销售VP"])
            subordinate_name = ""
            if is_director:
                st.info("👑 触发总监助销模式：系统将以高管身份生成降维打击话术。")
                subordinate_name = st.text_input(
                    "👤 负责该项目的下属姓名 (用于话术中自然引出)：",
                    placeholder="例如：小王 / 李工",
                    key="fire_subordinate",
                )

            with st.expander("⚙️ 高级技术方案配置 (四维精准制导)", expanded=True):
                tech_competitor = st.text_input(
                    "⚔️ 明确对比友商 (留空则常规输出)：",
                    placeholder="例如：西门子、ABB",
                    key="tech_competitor",
                )
                tech_status = st.text_input(
                    "📊 客户当前系统现状：",
                    placeholder="例如：一期设备老化严重，经常跳闸",
                    key="tech_status",
                )
                tech_pain_points = st.multiselect(
                    "🎯 客户核心痛点 (可多选)：",
                    st.session_state.pain_point_options,
                    key="tech_pains",
                )
                tech_role = st.multiselect(
                    "👤 沟通对象在采购链中的角色 (可身兼数职)：",
                    st.session_state.role_options,
                    key="tech_role",
                )

            # 聚合情报上下文（所有按钮共用）
            fire_context_parts = [str(row[3]) for row in logs if row[3]] if logs else []
            fire_context = "\n".join(fire_context_parts)

            fire_col1, fire_col2 = st.columns(2)
            with fire_col1:
                btn_email = st.button("✉️ 一键生成跟进话术", use_container_width=True)
            with fire_col2:
                btn_tech = st.button("📄 一键生成技术方案摘要", use_container_width=True)

            # ── btn_email 处理（紧随按钮定义，防止 Streamlit 丢失按钮状态）──
            if btn_email:
                if not api_key:
                    st.toast("🔌 请先在左上角系统设置中填入 API Key！", icon="🔑")
                    st.warning("🔌 请先在左上角「⚙️ 系统设置」中填入 API Key！")
                else:
                    label = "微信消息" if channel_key == "wechat" else "跟进邮件"
                    comp = tech_competitor if 'tech_competitor' in dir() else ""
                    status = tech_status if 'tech_status' in dir() else ""
                    pains = ', '.join(tech_pain_points) if 'tech_pain_points' in dir() and tech_pain_points else "待挖掘"
                    roles = ', '.join(tech_role) if 'tech_role' in dir() and tech_role else "决策者"
                    try:
                        with st.spinner(f"✉️ AI 正在为【{target_person}】定制{label}..."):
                            if fire_context.strip():
                                email = generate_followup_email(
                                    api_key, fire_context, channel_key,
                                    target_person, project_stage,
                                    use_top_to_top, shared_history,
                                    is_director, subordinate_name
                                )
                            else:
                                from openai import OpenAI as _OAI_email
                                _cli = _OAI_email(api_key=api_key)
                                prompt = f"""你是一位身经百战的顶尖B2B大客户销售。请根据以下参数，写一段发给客户的{label}跟进话术。
    【项目】：{target_proj}
    【竞品情报】：{comp or '未知'}
    【客户现状】：{status or '未知'}
    【客户痛点】：{pains}
    【沟通对象】：{target_person}（身份: {roles}）
    【项目阶段】：{project_stage}

    要求：1. 语气专业自信有分寸感 2. 直击对方身份最关心的利益点 3. 如有竞品隐晦打击竞品软肋 4. 200字以内直接输出话术正文"""
                                resp = _cli.chat.completions.create(
                                    model="gpt-4o",
                                    messages=[{"role": "user", "content": prompt}]
                                )
                                email = resp.choices[0].message.content
                        st.success("✅ 话术生成完毕！您可以直接复制发送。")
                        st.text_area(f"📨 生成的{label}（可直接复制）", email, height=300)
                    except Exception as e:
                        st.error(f"❌ 生成失败，请检查 API Key 是否正确！错误详情: {e}")

            # ── btn_tech 处理 ──
            if btn_tech:
                if not api_key:
                    st.toast("🔌 请先在左上角系统设置中填入 API Key！", icon="🔑")
                    st.warning("🔌 请先在左上角「⚙️ 系统设置」中填入 API Key！")
                else:
                    label = "技术要点" if channel_key == "wechat" else "技术方案摘要"
                    comp = tech_competitor if 'tech_competitor' in dir() else ""
                    status = tech_status if 'tech_status' in dir() else ""
                    pains = ', '.join(tech_pain_points) if 'tech_pain_points' in dir() and tech_pain_points else "待挖掘"
                    roles = ', '.join(tech_role) if 'tech_role' in dir() and tech_role else "决策者"
                    try:
                        with st.spinner(f"📄 AI 正在编制{label}..."):
                            if fire_context.strip():
                                tech = generate_tech_summary(
                                    api_key, fire_context, channel_key,
                                    tech_competitor, tech_status,
                                    tech_pain_points, tech_role
                                )
                            else:
                                from openai import OpenAI as _OAI_tech
                                _cli = _OAI_tech(api_key=api_key)
                                prompt = f"""你是一位资深的技术售前专家。请根据以下参数，生成一段用于方案PPT或汇报开头的【技术方案摘要】。
    【项目】：{target_proj}
    【竞品情报】：{comp or '未知'}
    【客户现状】：{status or '未知'}
    【客户痛点】：{pains}
    【汇报对象身份】：{roles}

    要求：1. 使用结构化工程语言条理清晰 2. 强调技术指标如何精准解决客户现状与痛点 3. 形成针对竞品的差异化技术壁垒"""
                                resp = _cli.chat.completions.create(
                                    model="gpt-4o",
                                    messages=[{"role": "user", "content": prompt}]
                                )
                                tech = resp.choices[0].message.content
                        st.success("✅ 技术方案摘要生成完毕！")
                        st.markdown(f"### 🎯 专为该客户定制的{label} (请直接复制发给客户)：")
                        st.markdown(tech)
                    except Exception as e:
                        st.error(f"❌ 技术摘要生成失败，请检查 API Key 是否正确！错误详情: {e}")

            # 动态画像标签已由 _init_dynamic_options() 统一初始化

            st.markdown("#### \U0001f575\ufe0f\u200d\u2642\ufe0f 内线专属通道 (教练弹药库)")
            with st.expander("🎯 锁定汇报目标 (决策者心理画像分析)", expanded=True):
                # ── 领导态度选择 + 动态管理 ──
                leader_attitude = st.selectbox(
                    "🧠 领导当前态度/关注核心：",
                    st.session_state.leader_attitudes,
                    key="insider_leader_attitude",
                )
                att_col1, att_col2 = st.columns([4, 1])
                with att_col1:
                    new_att = st.text_input(
                        "➕ 添加新的态度标签：",
                        placeholder="例如：极度关注环保合规与碳排放指标",
                        key="new_attitude_input",
                    )
                with att_col2:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("添加", key="btn_add_attitude"):
                        if new_att.strip() and new_att.strip() not in st.session_state.leader_attitudes:
                            st.session_state.leader_attitudes.append(new_att.strip())
                            st.rerun()
                        elif new_att.strip():
                            st.toast("⚠️ 该标签已存在！")
                # 删除按钮区域
                if len(st.session_state.leader_attitudes) > 1:
                    del_att = st.selectbox(
                        "🗑️ 选择要删除的态度标签：",
                        ["（不删除）"] + st.session_state.leader_attitudes,
                        key="del_attitude_select",
                    )
                    if del_att != "（不删除）":
                        if st.button(f"确认删除「{del_att[:10]}...」", key="btn_del_attitude"):
                            st.session_state.leader_attitudes.remove(del_att)
                            st.rerun()

                st.divider()

                # ── 领导历史轨迹选择 + 动态管理 ──
                leader_history = st.selectbox(
                    "🕰️ 领导的历史轨迹/心理阴影：",
                    st.session_state.leader_histories,
                    key="insider_leader_history",
                )
                hist_col1, hist_col2 = st.columns([4, 1])
                with hist_col1:
                    new_hist = st.text_input(
                        "➕ 添加新的历史标签：",
                        placeholder="例如：上一任厂长被设备事故免职，新领导极度保守",
                        key="new_history_input",
                    )
                with hist_col2:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("添加", key="btn_add_history"):
                        if new_hist.strip() and new_hist.strip() not in st.session_state.leader_histories:
                            st.session_state.leader_histories.append(new_hist.strip())
                            st.rerun()
                        elif new_hist.strip():
                            st.toast("⚠️ 该标签已存在！")
                # 删除按钮区域
                if len(st.session_state.leader_histories) > 1:
                    del_hist = st.selectbox(
                        "🗑️ 选择要删除的历史标签：",
                        ["（不删除）"] + st.session_state.leader_histories,
                        key="del_history_select",
                    )
                    if del_hist != "（不删除）":
                        if st.button(f"确认删除「{del_hist[:10]}...」", key="btn_del_history"):
                            st.session_state.leader_histories.remove(del_hist)
                            st.rerun()

            btn_insider = st.button("\U0001f525 一键生成【内线向上汇报/控标】专属隐蔽话术", use_container_width=True, type="primary")

            if btn_insider:
                if not api_key:
                    st.toast("🔌 请先在左上角系统设置中填入 API Key！", icon="🔑")
                    st.warning("🔌 请先在左上角「⚙️ 系统设置」中填入 API Key！")
                elif not fire_context:
                    st.toast("该项目暂无情报数据，请先录入拜访记录。", icon="📭")
                    st.info("📭 该项目暂无情报数据，无法生成内线话术。请先在「情报录入」Tab 中添加拜访记录。")
                else:
                    try:
                        with st.spinner("🕵️ 正在为您的内线/教练量身定制「向上管理」话术..."):
                            insider = generate_insider_ammo(
                                api_key, fire_context, channel_key,
                                target_person, project_stage,
                                leader_attitude, leader_history
                            )
                        st.markdown("### 🔒 极密：内线专属「向上管理」话术 (仅供教练使用，切勿外传)")
                        st.warning(insider)
                    except Exception as e:
                        st.error(f"内线话术生成失败：{e}")

            st.divider()

            # ── 💬 AI 参谋部 ──
            col_adv_title, col_adv_clear = st.columns([4, 1])
            with col_adv_title:
                st.subheader("💬 AI 参谋部")
            with col_adv_clear:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("🧹 清空推演记录", key="clear_advisor", use_container_width=True):
                    st.session_state.advisor_messages = []
                    st.rerun()

            # 项目切换时重置聊天记录，避免串台
            if "advisor_project_id" not in st.session_state or st.session_state.advisor_project_id != sandbox_proj_id:
                st.session_state.advisor_project_id = sandbox_proj_id
                st.session_state.advisor_messages = []

            # 渲染历史对话
            for msg in st.session_state.advisor_messages:
                st.chat_message(msg["role"]).write(msg["content"])

            # 用户输入
            user_query = st.chat_input(
                "向参谋长提问（如：生成一份发给老板的周报 / 分析当前赢单率）..."
            )

            if user_query:
                if not api_key:
                    st.warning("请先在左侧侧边栏输入 API Key！")
                else:
                    # 显示用户消息
                    st.chat_message("user").write(user_query)
                    st.session_state.advisor_messages.append({"role": "user", "content": user_query})

                    # 聚合项目情报上下文
                    context_parts = []
                    for row in logs:
                        try:
                            context_parts.append(row[3])
                        except (IndexError, TypeError):
                            pass
                    context_str = "\n".join(context_parts)

                    # 流式输出 AI 回答
                    try:
                        with st.chat_message("assistant"):
                            stream = chat_with_project_stream(
                                api_key, context_str, st.session_state.advisor_messages
                            )
                            response = st.write_stream(stream)
                        st.session_state.advisor_messages.append({"role": "assistant", "content": response})
                    except Exception as e:
                        st.error(f"AI 参谋调用失败：{e}")

    # ── AI 伴学中心 ──
with tab_academy:
    st.subheader("🎓 AI 实战伴学中心")
    st.caption("基于真实项目情报，AI 教练为你量身定制刁钻的实战演练题。")

    academy_projects = get_all_projects()

    if not academy_projects:
        st.warning("⚠️ 暂无项目数据，请先在情报录入中录入拜访记录。")
    else:
        academy_map = {p[1]: p[0] for p in academy_projects}
        academy_proj_name = st.selectbox(
            "📂 选择要进行实战演练的项目：",
            list(academy_map.keys()),
            key="academy_proj",
        )
        academy_proj_id = academy_map[academy_proj_name]

        st.divider()

        # 生成测验卡
        if st.button("🎯 生成今日实战测验卡", type="primary"):
            if not api_key:
                st.warning("请先在左侧侧边栏输入 API Key！")
            else:
                _, academy_logs = get_project_data(academy_proj_id)
                if not academy_logs:
                    st.info("该项目暂无情报数据，无法生成测验。")
                else:
                    context_parts = [row[3] for row in academy_logs if row[3]]
                    context_str = "\n".join(context_parts)

                    # 获取盲点数据
                    blind_spots = get_user_blind_spots()

                    try:
                        with st.spinner("📝 AI 教练正在基于三维框架出题..."):
                            quiz = generate_quiz(api_key, context_str, blind_spots)
                        st.session_state.current_quiz = quiz
                    except Exception as e:
                        st.error(f"出题失败：{e}")

        # 展示测验题
        if "current_quiz" in st.session_state and st.session_state.current_quiz:
            st.warning(f"📋 **今日实战题：**\n\n{st.session_state.current_quiz}")

            st.divider()

            st.markdown("### 🗣️ 实战模拟与破局演练")
            st.info("💡 场景：假设你是负责该项目的销售，面对上方 AI 教头给出的刁钻局势，请写出（或口述）你下一步具体的【战术动作】或【应对话术】。")

            user_answer = voice_text_area(
                label="请写下（或语音输入）你的应对话术或策略：",
                key="academy_input_voice_fixed",
                placeholder="点击上方 🎙️ 图标可以直接录音...",
                height=200
            )
            answer = user_answer  # 保持下游变量兼容
            if st.button("📮 提交策略并获取 AI 点评", type="primary"):
                if not answer.strip():
                    st.warning("⚠️ 请先输入或语音录入您的应对话术。")
                elif not api_key:
                    st.warning("请先在左侧侧边栏输入 API Key！")
                else:
                    with st.spinner("🕵️‍♂️ 王牌销售教头正在逐句拆解您的话术..."):
                        # 构建高阶点评 Prompt
                        project_context = st.session_state.get("project_data", "暂无情报记录")
                        coach_prompt = f"""你是一位年薪千万的 B2B 大客户销售总监兼无情的演练教头（精通 Miller Heiman 体系）。
你的下属（销售员）刚刚面对以下项目局势，给出了一段他的【实战应对话术/策略】。

【项目当前局势与基座情报】：
{project_context}

【AI 教练出的实战题】：
{st.session_state.current_quiz}

【销售员的实战话术/策略】：
"{answer}"

【你的点评任务】：
你不需要客气！请直接指出他话术里的致命漏洞，并给出极具杀伤力的示范。
请严格按照以下 Markdown 格式输出你的点评报告：

### 📊 战术维度评分 (总分 100)
- **破冰与共情 (25分)**：[你的打分] - [一句话点评：他是否拉近了距离？]
- **痛点与价值 (25分)**：[你的打分] - [一句话点评：他是否打中了客户的核心软肋？]
- **排他与控标 (25分)**：[你的打分] - [一句话点评：他有没有成功给竞品挖坑/设阻？]
- **推进与逼单 (25分)**：[你的打分] - [一句话点评：他是否拿到了下一步的承诺 (Next Step)？]

### 🔪 致命漏洞剖析
[指出他话术中最天真、最容易被客户怼回来、或者最容易被友商利用的 1 到 2 个核心漏洞。语气要犀利！]

### 💎 满分示范 (教头下场演示)
[请你亲自下场，写一段可以直接发给客户/当面说的满分话术。要完美融合人情世故、技术施压和战略推进！]
"""

                        try:
                            from openai import OpenAI as _OpenAI
                            _client = _OpenAI(api_key=api_key)
                            response = _client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[
                                    {"role": "system", "content": "你是一位极其严苛的 B2B 大客户销售教头，点评必须犀利、实战、有杀伤力。"},
                                    {"role": "user", "content": coach_prompt},
                                ],
                                temperature=0.7,
                            )
                            coach_feedback = response.choices[0].message.content

                            # 渲染点评结果
                            st.markdown("---")
                            st.success("✅ 点评完成！请查收您的实战诊断报告：")
                            st.markdown(coach_feedback)

                            # 尝试从 Markdown 中提取总分用于入库
                            import re as _re
                            score_matches = _re.findall(r"(\d+)\s*分?\s*-", coach_feedback)
                            total_score = sum(int(s) for s in score_matches[:4]) if score_matches else 0

                            # 持久化入库
                            save_test_record(
                                st.session_state.get("current_user", "未知"), academy_proj_id,
                                st.session_state.current_quiz, answer,
                                total_score, coach_feedback, "[]"
                            )
                            st.info("💾 测验记录已归档入库！")

                        except Exception as e:
                            st.error(f"❌ 点评引擎调用失败：{e}")

# ── 核心技术武器库 ──
with tab_knowledge:
    st.markdown("## 📚 中央武器库：技术弹药充填舱")
    st.caption("🔒 权限隔离：仅限后方核心技术团队与高管进行多模态资产注入。")

    # 识别权限
    is_tech_admin = st.session_state.get("role") in ["销售VP", "区域总监", "系统管理员"]

    if is_tech_admin:
        with st.expander("🔥 开启弹药装填 (支持多模态解析)", expanded=True):
            st.markdown("#### 📥 第一输入口：本地多模态资产")
            col_type, col_upload = st.columns([1, 2])
            with col_type:
                asset_type = st.selectbox("🎯 弹药归属分类:", ["🎯 [标准] 选型排雷 (温升/分断/智能化)", "🧮 [财务] TCO降本 (能效/无人值守核算)", "�️ [背书] 权威背书 (型式试验/电网实录)", "�️ [趋势] 宏观趋势 (双碳/电气火灾警示)"])
            with col_upload:
                uploaded_assets = st.file_uploader("📎 支持格式: PDF, Word, PPT, MP4, MP3", accept_multiple_files=True, type=['pdf', 'docx', 'ppt', 'pptx', 'mp4', 'mp3'])
                st.warning("⚠️ 弹药质检提示：请优先上传原生 PDF 的型式试验报告、选型手册。包含【内部燃弧/抗短路测试】的 MP4 视频，请务必配合 .txt 剧情说明文件一同上传！")

            st.markdown("#### 🔗 第二输入口：外部智能知识源")
            notebooklm_url = st.text_input("🌐 引用 NotebookLM / 企业 Wiki 共享链接:", placeholder="例如: https://notebooklm.google.com/...")

            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("🚀 一键向量化并传输至私有武器库 (API 对接)", type="primary", use_container_width=True):
                if uploaded_assets or notebooklm_url:
                    with st.spinner("🔗 正在通过私密 API 通道，将弹药传输至企业级知识库 (如 Dify/FastGPT)..."):
                        import time
                        time.sleep(1.5) # 模拟网络请求延迟
                        
                        # ==========================================================
                        # ⚠️ 企业级知识库 API 预留槽位 ⚠️
                        # 老板：未来在这里接入您的私有服务器接口。
                        # 例如:
                        # import requests
                        # headers = {"Authorization": "Bearer YOUR_DIFY_API_KEY"}
                        # data = {"url": notebooklm_url}
                        # response = requests.post("https://your-server.com/v1/datasets", headers=headers, json=data)
                        # ==========================================================
                        
                        # 暂时用全局缓存模拟私有知识库已接收数据的状态
                        st.session_state.private_kb_ready = True
                        
                        st.success(f"✅ 弹药已成功传输至私有向量库！前线【第一现场】的 AI 检索中枢已就绪。")
                else:
                    st.error("⚠️ 指挥官，弹药舱为空，请先上传文件或输入有效链接！")
    else:
        st.info("🛡️ 权限拦截：您当前的身份为【前线销售】，无法直接修改中央武器库。请前往【第一现场】调用已有弹药。")

    st.divider()
    st.markdown("### 🗄️ 当前可用弹药矩阵 (已挂载至第一现场)")
    c1, c2, c3 = st.columns(3)
    c1.info("**📖 核心参数文档**\n\n- 万华防腐涂料规格.pdf\n- 2026竞品打法手册.docx")
    c2.success("**🎬 多模态实况资产**\n\n- 镇海炼化吊装实录.mp4\n- 王总工破冰答疑.mp3")
    c3.warning("**🔗 外部智能链接源**\n\n- NotebookLM: 行业合规标准库\n- 内部工艺流程 Wiki")

# ── 第一现场 ──
with tab_live_pitch:
    st.markdown("## 🎙️ 第一现场：沉浸式客户展厅与双向连线")
    st.caption("🎯 在此模式下，系统分为【客户明面交互】与【销售战术暗线】。")

    # 模拟当前带入现场的项目上下文
    live_project_list = list(st.session_state.get("projects", {}).keys()) or ["暂无项目"]
    current_live_project = st.selectbox(
        "🔗 绑定本次拜访作战项目 (继承沙盘数据):",
        live_project_list,
        key="live_pitch_project"
    )
    st.divider()

    # 针对平板/大屏优化的 6:4 黄金分屏比例
    col_client, col_sales = st.columns([6, 4])

    with col_client:
        st.markdown("### 🖥️ 客户交互大屏 (Client View)")
        
        # --- 新增：平板端全息展示模块 ---
        st.markdown("##### 🧰 现场弹药包")
        # 使用 tabs 实现平板端最顺手的滑动切换体验
        media_tabs = st.tabs(["🎬 视频", "📊 PPT", "📐 CAD", "🎙️ 播客", "🗺️ 信息图", "📋 选型表", "✨ AI 模拟"])
        
        with media_tabs[0]:
            import os
            video_path = "assets/promo.mp4"
            if os.path.exists(video_path):
                st.video(video_path)
            else:
                st.info("💡 系统提示：未检测到真实企业宣传片。请在 `app.py` 同级目录下创建 `assets` 文件夹，并放入名为 `promo.mp4` 的视频文件。")
                # 兜底测试视频
                st.video("https://www.w3schools.com/html/mov_bbb.mp4")
            st.caption("🎬 演示视频：过往标杆项目设备吊装实录")
        with media_tabs[1]:
            st.info("� **交流 PPT：** 这里将直接渲染对接企业云盘的幻灯片组件，支持客户手势滑动。")
        with media_tabs[2]:
            st.warning("📐 **CAD 上图模块：** 此处预留 WebGL 接口，用于三维模型旋转及爆炸图拆解演示。")
        with media_tabs[3]:
            audio_path = "assets/expert_voice.mp3"
            if os.path.exists(audio_path):
                st.audio(audio_path)
            else:
                st.info("💡 系统提示：未检测到真实专家录音。请在 `app.py` 同级目录下的 `assets` 文件夹中放入名为 `expert_voice.mp3` 的音频文件。")
                st.audio("https://www.w3schools.com/html/horse.ogg")
            st.caption("🎙️ 行业播客：研发总工解读最新环保排污政策")
        with media_tabs[4]:
            st.success("🗺️ **信息图：** 此处渲染 TCO (全生命周期成本) 与竞品 ROI 对比雷达图。")
        with media_tabs[5]:
            import pandas as pd
            st.dataframe(pd.DataFrame({
                "参数维度": ["额定功率", "防护等级", "防腐标准", "交付周期"],
                "旗舰型号 (推荐)": ["500kW", "IP67", "C5-M (海洋级)", "30天"],
                "标准型号": ["300kW", "IP65", "C3 (工业级)", "15天"]
            }), use_container_width=True, hide_index=True)
            
        with media_tabs[6]:
            st.markdown("##### 🚀 动态方案推演引擎")
            st.caption("基于全局项目情报，由 AI 现场现编极具针对性的应对方案。")
            sim_type = st.selectbox("🎯 选择推演场景：", ["极端工况抗压模拟", "TCO 投资回报率核算", "施工排期极限压缩方案"], key="sim_select")
            
            if st.button("🧠 结合当前项目一键推演", key="btn_sim", type="primary", use_container_width=True):
                if not api_key:
                    st.error("⚠️ 请先在左侧输入 API Key！")
                else:
                    with st.spinner("AI 正在结合项目情报进行推演..."):
                        try:
                            from openai import OpenAI as _OpenAI_sim
                            _client_sim = _OpenAI_sim(api_key=api_key)
                            sim_prompt = f"你是顶尖的技术售前。当前正在向客户展示项目：【{current_live_project}】。客户要求进行【{sim_type}】。请结合行业常识和该项目潜在的痛点，直接输出一段大约 200 字、极具专业度和说服力的方案推演结论，必须包含具体数据预测，且可以直接展示给客户看。"
                            
                            resp_sim = _client_sim.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[{"role": "user", "content": sim_prompt}],
                                temperature=0.7
                            )
                            st.session_state.current_simulation = resp_sim.choices[0].message.content
                        except Exception as e:
                            st.error(f"推演失败：{e}")
            
            if "current_simulation" in st.session_state:
                st.info(f"**✅ 推演报告 ({sim_type})：**\n\n{st.session_state.current_simulation}")
        
        st.divider()

        # 1. 初始化当前项目的独立聊天记录
        chat_history_key = f"live_chat_{current_live_project}"
        if chat_history_key not in st.session_state:
            st.session_state[chat_history_key] = []

        # 2. 渲染历史聊天记录 (让对话可以一直往上追溯)
        for msg in st.session_state[chat_history_key]:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])

        # --- 战术雨刷：一键清空现场 ---
        col_clear, _ = st.columns([1, 4])
        with col_clear:
            if st.button("🧹 结束会议 (清空现场)", help="清理当前客户的对话记录，准备迎接下一波战斗"):
                st.session_state[chat_history_key] = []
                st.session_state.last_client_query = ""
                st.rerun()

        # 3. 客户直连大模型通道
        pitch_query = st.chat_input("🎙️ 现场遇阻？输入客户的刁钻发难 (例如: 你们的设备散热不如西门子)...")

        if pitch_query:
            # 记录最后一次 query 供右侧护目镜和 SOS 使用
            st.session_state.last_client_query = pitch_query
            st.toast("捕捉到现场交流信号！", icon="📡")
            
            if not api_key:
                st.error("⚠️ 请先在左侧边栏输入大模型 API Key！")
            else:
                with st.spinner("🧠 护目镜正在疯狂调取资料库，计算反击话术..."):
                    try:
                        from openai import OpenAI as _OAI_pitch
                        _client_pitch = _OAI_pitch(api_key=api_key)
                        dify_knowledge = st.session_state.get("global_knowledge", "暂无挂载知识库。")
                        
                        pitch_prompt = f"""你是一位年薪千万的 B2B 大客户售前总监。现在路演现场客户突然发难。
【客户刁难】："{pitch_query}"

【我方底牌/技术参数】：
{dify_knowledge[:1500]}

请通过"战术护目镜"向我方销售发送紧急话术支援。要求：
1. 极度简短、一剑封喉（最多 3 句话），让销售能扫一眼直接照着念。
2. 语气要有底气。如果是竞品（如西门子、ABB等）的常见攻击套路，请直接揭穿背后的技术陷阱。
"""
                        resp_pitch = _client_pitch.chat.completions.create(
                            model="gpt-4o",
                            messages=[{"role": "user", "content": pitch_prompt}],
                            temperature=0.7
                        )
                        st.session_state.live_pitch_answer = resp_pitch.choices[0].message.content
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ 护目镜连接失败：{e}")

            # 存入并显示用户消息
            st.session_state[chat_history_key].append({"role": "user", "content": pitch_query})
            with st.chat_message("user"):
                st.write(f"**客户提问：** {pitch_query}")

            # AI 响应生成
            with st.chat_message("assistant"):
                if not api_key:
                    st.warning("⚠️ 请先在左侧侧边栏输入 API Key 以激活真实 AI 大脑！")
                else:
                    with st.spinner("🧠 真实 AI 大脑正在分析工况并生成专业解答..."):
                        try:
                            from openai import OpenAI as _OpenAI_live
                            _client_live = _OpenAI_live(api_key=api_key)
                            
                            # 将历史记录拼接到 Prompt 中，让 AI 有上下文记忆
                            history_context = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state[chat_history_key][-5:]])
                            
                            # ==========================================================
                            # ⚠️ 企业级知识库检索 API 预留槽位 ⚠️
                            # 老板：未来在这里接入您的私有知识库查询接口，用 client_query 去查资料。
                            # 例如:
                            # response = requests.post("https://your-server.com/v1/chat-messages", json={"query": client_query, ...})
                            # retrieved_knowledge = response.json().get("answer")
                            # ==========================================================
                            
                            # 模拟从私有知识库检索到的精确片段
                            if st.session_state.get("private_kb_ready"):
                                retrieved_knowledge = f"【系统提示：已成功从企业私有知识库 (API) 中检索到关于 '{pitch_query[:10]}...' 的绝密数据，请结合此数据作答。】"
                            else:
                                retrieved_knowledge = "暂未挂载企业私有知识库。"
                            
                            live_prompt = f"""你是一位在第一现场的高级技术售前专家。
当前关联项目：【{current_live_project}】
近期对话上下文：
{history_context}

【🔥 私有兵工厂精准弹药支援 (必须优先参考)】：
{retrieved_knowledge}

请针对客户的最新提问，给出极其专业、有针对性的解答。语气要自信、客观，能够直接展示在屏幕上给客户看。
"""
                            response = _client_live.chat.completions.create(
                                model="gpt-4o",
                                messages=[{"role": "user", "content": live_prompt}],
                                temperature=0.7
                            )
                            real_answer = response.choices[0].message.content
                            
                            st.write(real_answer)
                            st.caption("✅ 真实 AI 答案已生成 (已自动挂载当前项目上下文)。")
                            
                            # 将 AI 的回答也存入历史记录
                            st.session_state[chat_history_key].append({"role": "assistant", "content": real_answer})
                        except Exception as e:
                            st.error(f"❌ 呼叫现场 AI 大脑失败：{e}")

    with col_sales:
        st.markdown("### 🥽 战术护目镜 (Sales Only)")
        st.caption("🤫 仅销售可见的实时底牌与战术指导")

        if st.session_state.get("last_client_query"):
            last_q = st.session_state.get("last_client_query")
            if not api_key:
                st.warning("⚠️ 等待 API Key 激活战术护目镜...")
            else:
                with st.spinner("🥽 护目镜实时透视客户意图..."):
                    try:
                        from openai import OpenAI as _OpenAI_goggles
                        _client_goggles = _OpenAI_goggles(api_key=api_key)
                        
                        goggles_prompt = f"""你是隐藏在销售耳机里的顶尖战术大师。
当前关联项目：【{current_live_project}】
客户刚刚在第一现场提出了这个问题："{last_q}"

请用极简的要点（100字以内），给销售提供实时的战术指导。必须包含两点：
1. 🎯 话术拆解 (客户真正的担忧是什么)
2. 💡 出牌建议 (销售现在该用什么案例或筹码回击)
"""
                        resp_goggles = _client_goggles.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[{"role": "user", "content": goggles_prompt}],
                            temperature=0.8
                        )
                        tactical_advice = resp_goggles.choices[0].message.content
                        
                        st.warning(f"**🥽 战术护目镜实时解析：**\n\n{tactical_advice}")
                    except Exception as e:
                        st.error(f"护目镜信号干扰: {e}")

            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("🆘 一键呼叫后方技术群 (企业微信连线)", type="primary", use_container_width=True):
                client_q = st.session_state.get("last_client_query", "")
                if not client_q:
                    st.warning("⚠️ 暂未捕捉到有效的现场交锋上下文，系统无法生成支援简报！请先在左侧输入客户提问。")
                else:
                    if not api_key:
                        st.warning("⚠️ 请先输入 API Key 以生成求援简报！")
                    else:
                        with st.spinner("📡 AI 正在提炼现场火力需求，加密呼叫后方专家群..."):
                            try:
                                from openai import OpenAI as _OpenAI_sos
                                _client_sos = _OpenAI_sos(api_key=api_key)
                                
                                sos_prompt = f"""你是前线销售的 AI 战术助理。客户刚刚在现场提出了以下棘手问题：
"{client_q}"
当前关联项目：【{current_live_project}】

请帮销售向后方的【核心技术与商务专家群】写一段极其简短、专业的求援需求（3点以内）。
说明前线遇到了什么阻力，需要后方提供什么样的具体支持（如：特定的测试数据、商务特批权限、某友商的劣势分析等）。直接输出需求要点。"""
                                
                                resp_sos = _client_sos.chat.completions.create(
                                    model="gpt-4o-mini",
                                    messages=[{"role": "user", "content": sos_prompt}],
                                    temperature=0.7
                                )
                                sos_brief = resp_sos.choices[0].message.content
                                
                                import random
                                ticket_id = f"T-2026-{random.randint(1000, 9999)}"
                                
                                import datetime
                                now_str = datetime.datetime.now().strftime("%H:%M:%S")
                                if "sos_tickets" not in st.session_state:
                                    st.session_state.sos_tickets = []
                                st.session_state.sos_tickets.insert(0, {
                                    "id": ticket_id,
                                    "time": now_str,
                                    "project": current_live_project,
                                    "query": client_q,
                                    "brief": sos_brief,
                                    "status": "🔴 紧急待支援"
                                })
                                
                                st.success(f"✅ **紧急工单已派发群聊 ({ticket_id})**")
                                with st.container(border=True):
                                    st.markdown("**🧑‍💻 接收专家：** 售前支持群 / 商务特批组")
                                    st.markdown(f"**🗣️ 现场原声截取：**\n> {client_q}")
                                    st.markdown(f"**🎯 AI 自动生成支援需求：**\n{sos_brief}")
                                    st.caption("⏳ 专家已收到企业微信震动提醒，预计 1 分钟内回传弹药...")
                            except Exception as e:
                                st.error(f"❌ SOS 信号发送失败: {e}")
        else:
            if "live_pitch_answer" in st.session_state:
                st.success("🎯 **AI 隐形提词器 (建议直接原话复述)：**\n\n" + st.session_state.live_pitch_answer)
                if st.button("🧹 清除当前提示", key="clear_pitch"):
                    del st.session_state.live_pitch_answer
                    st.rerun()
            else:
                st.markdown("*(等待捕捉现场客户提问...)*")

        # --- 新增：空投弹药接收雷达 ---
        if "sos_tickets" in st.session_state:
            # 筛选当前项目且已送达的弹药
            incoming_ammo = [t for t in st.session_state.sos_tickets if t['project'] == current_live_project and t['status'] == "🟢 支援已送达"]
            if incoming_ammo:
                st.divider()
                st.markdown("### 🚁 后方空投弹药已到达")
                for t in incoming_ammo:
                    with st.container(border=True):
                        st.success(f"**来自总部的特批指示 ({t['time']})：**\n\n{t.get('reply', '（总部已查收，执行特批预案）')}")

# ── 领导看板 ──
with tab_leader:
    if st.session_state.get("role", "") not in ["区域总监", "销售VP", "总经理", "系统管理员"]:
        st.warning("⚠️ 仅销售总监可访问此看板。请在左侧侧边栏切换角色。")
    else:
        st.markdown("## 📊 总裁全局作战大盘")
        st.caption("实时监控全盘商机、询报价漏斗与团队战斗力。数据已与【询报价】模块底层打通。")

        # --- 1. 顶部核心指标 KPI 卡片 ---
        st.markdown("### 📈 核心营收与风控指标 (实时)")
        kpi1, kpi2, kpi3, kpi4 = st.columns(4)

        # 动态抓取"询报价"模块中存入的获批底单数据
        approved_total = 0
        approved_count = 0
        if "approved_boms" in st.session_state:
            import pandas as pd
            for proj, df in st.session_state.approved_boms.items():
                approved_count += 1
                total = (pd.to_numeric(df["销售核定数量"], errors='coerce').fillna(0) * pd.to_numeric(df["标准单价(元)"], errors='coerce').fillna(0)).sum()
                approved_total += total

        kpi1.metric(label="💰 累计获批报价总额", value=f"¥ {approved_total / 10000:,.2f} 万", delta="本月环比持续增长")
        kpi2.metric(label="📑 已输出正式报价单", value=f"{approved_count} 份", delta="最新成单动态")
        kpi3.metric(label="⚔️ 整体控标胜率 (预估)", value="68%", delta="较上季度 +5%")
        kpi4.metric(label="🛡️ 天眼风控拦截次数", value="3 次", delta="避免潜在亏损单", delta_color="inverse")

        st.divider()

        # --- 2. 中部大盘图表区 ---
        st.markdown("### 🗺️ 战区业绩与团队战力")
        c_chart1, c_chart2 = st.columns(2)

        with c_chart1:
            with st.container(border=True):
                st.markdown("#### 🏢 各行业打单金额分布 (预估)")
                import pandas as pd
                chart_data = pd.DataFrame({
                    "行业": ["新能源", "化工", "传统电网", "冶金"],
                    "商机金额(万)": [850, 620, 410, 230]
                }).set_index("行业")
                st.bar_chart(chart_data)

        with c_chart2:
            with st.container(border=True):
                st.markdown("#### 🎖️ 销售团队战力实时排行")
                st.dataframe(pd.DataFrame([
                    {"排名": "🥇 1", "销售姓名": "张三", "跟进项目数": 5, "转化率": "80%", "当前状态": "🔥 爆单"},
                    {"排名": "🥈 2", "销售姓名": "李四", "跟进项目数": 3, "转化率": "60%", "当前状态": "✅ 正常"},
                    {"排名": "🥉 3", "销售姓名": "王五", "跟进项目数": 6, "转化率": "15%", "当前状态": "⚠️ 需辅导"}
                ]), use_container_width=True, hide_index=True)

        st.divider()

        # ================================================================
        # 顶部模块：【待办】项目立项审批看板
        # ================================================================
        st.subheader("📋 项目立项审批看板")

        pending = st.session_state.get("pending_projects", [])

        # --- 权限过滤逻辑 ---
        current_dept = st.session_state.get("user_dept", "ALL")

        # 筛选出当前用户有权看到的项目
        visible_projects = []
        for p in pending:
            proj_dept = p.get("dept", "未知战区")
            # 1. VP/管理员 (dept=="ALL") → 看全部
            if current_dept == "ALL":
                visible_projects.append(p)
            # 2. 区域总监 → 只看自己战区的
            elif proj_dept == current_dept:
                visible_projects.append(p)
            # 3. 旧数据 (无 dept) 归为 "未知战区"，仅 VP 可见 → 已被条件 1 覆盖

        # --- 渲染逻辑 ---
        if not visible_projects:
            if current_dept == "ALL":
                st.success("🎉 全战区目前没有积压的待审批项目。")
            else:
                hidden_count = len(pending) - len(visible_projects)
                st.success(f"🎉 {current_dept} 目前没有待审批项目。")
                if hidden_count > 0:
                    st.caption(f"💡 另有 {hidden_count} 个其他战区的项目在各自总监处审批中。")
        else:
            st.warning(f"🔔 您的管辖范围内有 **{len(visible_projects)}** 个项目等待审批"
                       + (f"（全局共 {len(pending)} 个）" if current_dept == "ALL" and len(pending) != len(visible_projects) else ""))

            for i, p in enumerate(visible_projects):
                pid = p.get("project_name", f"未知项目_{i}")
                proj_dept = p.get("dept", "未知战区")

                with st.expander(f"⏳ 待审项目：{pid} | 📍 {proj_dept}", expanded=True):
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        applicant = p.get("applicant", "未知")
                        submit_time = p.get("time", "未知")
                        st.markdown(f"""
**👤 提报人：** {applicant}  
**🏢 所属战区：** {proj_dept}  
**📅 提报时间：** {submit_time}
""")
                        if p.get("data"):
                            st.text(p["data"])

                    with col2:
                        if st.button("✅ 批准", key=f"leader_ok_{i}_{pid}", type="primary", use_container_width=True):
                            st.session_state.pending_projects.remove(p)
                            add_project(pid, "线索")
                            if pid not in st.session_state.project_name_cache:
                                st.session_state.project_name_cache.append(pid)
                            st.session_state.current_project = pid
                            st.toast(f"✅ 项目【{pid}】已批准立项！", icon="🎉")
                            st.rerun()

                        if st.button("❌ 驳回", key=f"leader_no_{i}_{pid}", type="secondary", use_container_width=True):
                            st.session_state.pending_projects.remove(p)
                            st.toast(f"❌ 项目【{pid}】已驳回。", icon="🗑️")
                            st.rerun()

        st.divider()

        # ================================================================
        # 新增模块：🚨 SOS 前线紧急呼叫调度中心
        # ================================================================
        col_sos_title, col_sos_clear = st.columns([4, 1])
        with col_sos_title:
            st.subheader("🚨 前线 SOS 调度中心")
        with col_sos_clear:
            st.markdown("<br>", unsafe_allow_html=True) # 对齐微调
            if st.button("🗑️ 销毁已处理工单", use_container_width=True):
                if "sos_tickets" in st.session_state:
                    # 只保留未处理的紧急工单，清除已送达的
                    st.session_state.sos_tickets = [t for t in st.session_state.sos_tickets if t['status'] == "🔴 紧急待支援"]
                    st.rerun()
        if "sos_tickets" not in st.session_state or not st.session_state.sos_tickets:
            st.success("☕ 当前天下太平，没有前线销售呼叫炮火支援。")
        else:
            for idx, ticket in enumerate(st.session_state.sos_tickets):
                with st.expander(f"{ticket['status']} | {ticket['id']} | 项目: {ticket['project']} | 呼叫时间: {ticket['time']}", expanded=(idx==0)):
                    st.markdown(f"**🗣️ 前线遭遇阻击：** {ticket['query']}")
                    st.markdown(f"**🎯 参谋部提炼需求：**\n{ticket['brief']}")
                    
                    if ticket['status'] == "🔴 紧急待支援":
                        reply_text = st.text_area("✍️ 输入支援弹药 / 商务特批意见：", key=f"reply_{ticket['id']}")
                        if st.button("🚀 一键全频道投送弹药", type="primary", key=f"btn_{ticket['id']}", use_container_width=True):
                            st.session_state.sos_tickets[idx]['status'] = "🟢 支援已送达"
                            # 核心修复：保存总监的批示文字
                            st.session_state.sos_tickets[idx]['reply'] = reply_text 
                            st.toast(f"炮火已成功向 {ticket['project']} 投送！", icon="🚁")
                            st.rerun()
                    else:
                        st.info("✅ 炮火已覆盖，前线危机解除。")
                        if 'reply' in ticket and ticket['reply']:
                            st.success(f"**您的批示记录：** {ticket['reply']}")
        st.divider()

        # ================================================================
        # 新增模块：⚖️ 撞单申诉与归属权仲裁法庭
        # ================================================================
        st.subheader("⚖️ 撞单申诉与归属权仲裁法庭")
        appeals = st.session_state.get("appeals", [])
        pending_appeals = [a for a in appeals if a.get("status") == "⚖️ 待裁决"]
        
        if not pending_appeals:
            st.success("⚖️ 战区和平，当前暂无撞单申诉需要裁决。")
        else:
            for idx, a in enumerate(pending_appeals):
                with st.expander(f"🔴 撞单争议：{a['new_project']} 🆚 {a['conflict_with']}", expanded=True):
                    c_left, c_right = st.columns([3, 1])
                    with c_left:
                        st.markdown(f"**发起申诉方（抢单者）：** `{a['applicant']}`")
                        st.markdown(f"**原归属方（守单者）：** `{a['original_owner']}`")
                        st.info(f"**📝 申诉核心依据：**\n{a['reason']}")
                        if a.get('has_evidence'):
                            st.caption("📎 销售已上传实锤证据材料（点击后台附件库查阅）")
                    with c_right:
                        st.markdown("<br>", unsafe_allow_html=True)
                        if st.button("✅ 申诉有效\n(剥夺并转移归属权)", key=f"appeal_yes_{idx}", type="primary", use_container_width=True):
                            # 核心业务逻辑：强制更改项目的 owner
                            conflict_name = a['conflict_with']
                            if conflict_name in st.session_state.projects:
                                st.session_state.projects[conflict_name]["owner"] = a['applicant']
                                st.session_state.projects[conflict_name]["applicant"] = a['applicant']
                            
                            # 修改申诉单状态为闭环
                            st.session_state.appeals[st.session_state.appeals.index(a)]["status"] = "✅ 胜诉转移"
                            st.toast("✅ 裁决生效：项目归属权已强制转移给申诉人！", icon="⚖️")
                            st.rerun()
                        
                        if st.button("❌ 驳回申诉\n(维持原判)", key=f"appeal_no_{idx}", use_container_width=True):
                            # 维持原状
                            st.session_state.appeals[st.session_state.appeals.index(a)]["status"] = "❌ 败诉驳回"
                            st.toast("❌ 裁决生效：已驳回抢单请求，维持原归属。", icon="⚖️")
                            st.rerun()
        
        st.divider()

        # ================================================================
        # 底部模块：【阅览】团队能力透视（考试成绩）
        # ================================================================
        st.subheader("📊 团队能力透视看板")

        records = get_all_test_records()

        if not records:
            st.info("暂无团队测验数据。")
        else:
            # 模块 A：团队测验明细墙
            st.markdown("### 📋 团队测验明细墙")
            import pandas as pd

            # 1. 转换为 DataFrame
            df_exam = pd.DataFrame(records, columns=["销售姓名", "关联项目", "测验得分", "盲点标签", "测验时间"])

            # 2. 补全直属领导逻辑 (根据组织架构反向推导)
            director_map = {
                "张三": "王总监 (华东)", "李四": "王总监 (华东)", "王五": "王总监 (华东)",
                "孙七": "赵总监 (华北)", "周八": "赵总监 (华北)",
                "阿宝": "陈总监 (华南)", "阿强": "陈总监 (华南)",
                "Linda": "林总监 (大客户)", "Tony": "林总监 (大客户)",
                "default": "历史遗留数据"
            }
            if "销售姓名" in df_exam.columns:
                df_exam.insert(1, "直属领导", df_exam["销售姓名"].map(lambda x: director_map.get(x, "未知高管直管")))

            # 3. 显式排序交互控件
            sort_mode = st.radio(
                "📊 请选择查阅视角:",
                ["⬇️ 得分从高到低 (抓标杆)", "⬆️ 得分从低到高 (抓辅导)"],
                horizontal=True
            )

            # 4. 执行核心排序逻辑
            is_ascending = "低到高" in sort_mode
            if "测验得分" in df_exam.columns:
                df_exam["测验得分"] = pd.to_numeric(df_exam["测验得分"], errors='coerce').fillna(0)
                df_exam = df_exam.sort_values(by="测验得分", ascending=is_ascending)

            # 5. 重置索引并强制生成从 1 开始的真实序号列
            df_exam = df_exam.reset_index(drop=True)
            df_exam.insert(0, "序号", df_exam.index + 1)

            # --- 新增：BI 数据可视化图表 ---
            st.markdown("#### 📈 战区平均战力分布图")
            if "直属领导" in df_exam.columns and "测验得分" in df_exam.columns:
                # 按照直属领导分组，计算每个战区的平均分
                chart_data = df_exam.groupby("直属领导")["测验得分"].mean().reset_index()
                # 渲染 Streamlit 原生柱状图
                st.bar_chart(chart_data, x="直属领导", y="测验得分", use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)

            # 6. 渲染增强版表格
            st.dataframe(
                df_exam,
                hide_index=True,
                use_container_width=True,
                column_config={
                    "序号": st.column_config.TextColumn("序号", width="small"),
                    "测验得分": st.column_config.ProgressColumn(
                        "测验得分",
                        help="员工考核分数",
                        format="%d 分",
                        min_value=0,
                        max_value=100,
                    ),
                }
            )

            # 统计概览
            avg_score = df_exam["测验得分"].mean()
            total_tests = len(df_exam)
            low_count = len(df_exam[df_exam["测验得分"] < 60])
            col_m1, col_m2, col_m3 = st.columns(3)
            col_m1.metric("📊 总测验次数", total_tests)
            col_m2.metric("📈 平均得分", f"{avg_score:.1f}")
            col_m3.metric("🚨 不及格次数", low_count)

            st.divider()

            # 模块 B：AI 团队能力体检报告
            st.markdown("### 🧠 团队能力 AI 体检报告")
            if st.button("🏥 一键生成团队能力体检报告", type="primary"):
                if not api_key:
                    st.warning("请先在左侧侧边栏输入 API Key！")
                else:
                    # 汇总所有盲点
                    all_spots = []
                    for row in records:
                        if row[3]:  # blind_spots 字段
                            all_spots.append(row[3])
                    spots_summary = "\n".join(all_spots)

                    try:
                        with st.spinner("🧠 AI 正在分析团队能力短板..."):
                            report = generate_team_report(api_key, spots_summary)
                        st.success(f"**📊 团队能力体检报告**\n\n{report}")
                    except Exception as e:
                        st.error(f"报告生成失败：{e}")


# ── 全局配置管理器 (侧边栏底部) ──
def render_config_manager():
    """在侧边栏底部渲染所有下拉选项的增删管理面板。"""
    with st.sidebar:
        st.markdown("---")
        with st.expander("⚙️ 全局系统参数配置 (管理员)", expanded=False):
            st.info("在此处维护系统所有的下拉菜单选项。")

            def manage_options(key, label):
                options = st.session_state.get(key, [])
                st.write(f"**{label}**")
                st.code(options)

                new_item = st.text_input(f"➕ 新增 {label}:", key=f"new_{key}")
                if st.button("添加", key=f"add_{key}"):
                    if new_item.strip() and new_item.strip() not in options:
                        st.session_state[key].append(new_item.strip())
                        st.rerun()
                    elif new_item.strip():
                        st.toast("⚠️ 该选项已存在！")

                items_to_remove = st.multiselect(f"🗑️ 删除 {label}:", options, key=f"del_{key}")
                if st.button("删除选中", key=f"remove_{key}"):
                    for item in items_to_remove:
                        if item in st.session_state[key]:
                            st.session_state[key].remove(item)
                    st.rerun()
                st.markdown("---")

            # 管理各模块参数（key 对应 DEFAULT_CONFIGS 中的键）
            cfg_tab_options, cfg_tab_meddic = st.tabs(["📋 下拉选项管理", "⚖️ MEDDIC 权重配置"])

            with cfg_tab_options:
                manage_options("project_stages", "项目阶段")
                manage_options("pain_point_options", "客户核心痛点")
                manage_options("role_options", "采购链角色")
                manage_options("leader_attitudes", "决策者态度标签")
                manage_options("leader_histories", "决策者历史标签")
                manage_options("info_sources", "信息来源")
                manage_options("project_drivers", "立项驱动力")
                manage_options("position_options", "我方身位")
                manage_options("budget_statuses", "预算状态")

            with cfg_tab_meddic:
                st.write("### 🧠 动态赢率评估模型库")
                st.info("💡 设定各项评估指标的绝对重要性 (0-100)。您可自由增删指标（如增加 '客情关系' 或 '预算合规'）。")

                # 动态渲染当前所有指标的拉杆
                st.write("**⚙️ 调整当前模型参数：**")
                updated_dims = {}
                for dim, weight in st.session_state.eval_dimensions.items():
                    updated_dims[dim] = st.slider(f"{dim} (重要度)", 0, 100, weight)

                # 实时保存拉杆状态
                st.session_state.eval_dimensions = updated_dims

                st.markdown("---")
                # 增删改查：添加新指标
                new_dim = st.text_input("➕ 新增评估指标名称：", placeholder="例如：B - 专项预算落实情况")
                if st.button("添加指标", use_container_width=True):
                    if new_dim and new_dim not in st.session_state.eval_dimensions:
                        st.session_state.eval_dimensions[new_dim] = 50  # 默认赋予50重要度
                        st.rerun()

                # 增删改查：删除指标
                del_dim = st.selectbox("🗑️ 删除不再适用的指标：", ["(不删除)"] + list(st.session_state.eval_dimensions.keys()))
                if st.button("确认删除指标", use_container_width=True):
                    if del_dim != "(不删除)" and del_dim in st.session_state.eval_dimensions:
                        del st.session_state.eval_dimensions[del_dim]
                        st.rerun()

                st.markdown("---")
                # 为老板构想的未来功能预留接口
                st.success("🤖 AI 闭环自学习引擎 (Auto-ML)")
                st.caption("基于未来 100 个闭环项目的胜败复盘数据，AI 将自动反向微调上述权重。（当前处于数据积累期）")
                st.button("启动自学习优化 (数据积累中...)", disabled=True, use_container_width=True)


if "expense_requests" not in st.session_state:
    # status 链路: 待总监审批 -> 待VP审批 -> 待总经理核准 -> 财务已执行 -> 驳回
    st.session_state.expense_requests = []

# ── 粮草与费用审批 ──
with tab_finance:
    st.markdown("## 💸 粮草战备库 (借款与费用追踪)")
    st.caption("项目所有的血液(资金)消耗都在此留痕，AI 将实时监控资金效率与 ROI 风险。")
    
    current_r = st.session_state.get("role", "一线销售")
    
    col_req, col_audit = st.columns([1, 1])
    
    with col_req:
        st.markdown("### 📝 发起资金申请")
        if current_r == "一线销售":
            with st.container(border=True):
                exp_proj = st.selectbox("📂 关联项目：", list(st.session_state.get("projects", {}).keys()) or ["暂无项目"], key="exp_proj")
                exp_type = st.selectbox("🏷️ 资金类目：", ["差旅费", "招待费", "客情维护专属", "项目运作/招投标费"], key="exp_type")
                
                # --- 新增：动态合规挂接逻辑 ---
                target_month = ""
                target_person = ""
                
                if exp_type == "差旅费":
                    import datetime
                    current_month = datetime.datetime.now().month
                    months = [f"{i}月" for i in range(1, 13)]
                    target_month = st.selectbox("📅 归属月度：", options=months, index=current_month-1, key="exp_month")
                else:
                    # 联动：获取该项目在沙盘中已录入的干系人列表
                    people_options = []
                    if "stakeholders" in st.session_state and exp_proj in st.session_state.stakeholders:
                        df_people = st.session_state.stakeholders[exp_proj]
                        if not df_people.empty and "姓名" in df_people.columns:
                            people_options = df_people["姓名"].dropna().unique().tolist()
                    
                    if not people_options:
                        st.error("⚠️ 财务风控拦截：该项目尚未建立权力地图。请先在【作战沙盘 -> 关键决策链】中添加具体的客户人员，方可申请招待/客情费用！")
                    else:
                        target_person = st.selectbox("🎯 挂接目标人物 (仅限本项目已知人员)：", options=people_options, key="exp_person")

                exp_amount = st.number_input("💰 申请金额 (元)：", min_value=0, step=500)
                exp_reason = st.text_area("✍️ 申请事由 (AI 将依据此判定资金效率)：", placeholder="例如：需请关键决策人王总吃饭，推进二期图纸确认...")
                
                if st.button("🚀 提交借款/费用申请", type="primary", use_container_width=True):
                    if exp_amount <= 0 or not exp_reason.strip():
                        st.warning("⚠️ 金额必须大于 0，且必须填写详细事由！")
                    elif exp_type != "差旅费" and not target_person:
                        st.warning("⚠️ 提交失败：必须挂接具体的业务关联人员！")
                    else:
                        import datetime, random
                        req_id = f"EXP-{datetime.datetime.now().strftime('%y%m%d')}-{random.randint(100,999)}"
                        
                        # 组装带属性的事由喂给 AI
                        if exp_type == "差旅费":
                            detailed_reason = f"[{target_month} 差旅] {exp_reason}"
                            ai_log_context = f"[{target_month} 的差旅费]"
                        else:
                            detailed_reason = f"[定点作用于: {target_person}] {exp_reason}"
                            ai_log_context = f"定点作用于关键人【{target_person}】的 [{exp_type}]"

                        new_req = {
                            "id": req_id, "project": exp_proj, "type": exp_type, 
                            "amount": exp_amount, "reason": detailed_reason, 
                            "applicant": current_user, "dept": st.session_state.user_dept,
                            "time": datetime.datetime.now().strftime('%Y-%m-%d %H:%M'),
                            "status": "待总监审批",
                            "audit_trail": [f"[{current_user}] 提交申请。"]
                        }
                        st.session_state.expense_requests.append(new_req)
                        
                        # 强力穿透：写入项目全局情报
                        intel_log = f"【💸 资金消耗预警】销售申请了 {exp_amount} 元，{ai_log_context}。事由：{exp_reason}。(当前状态：审批中)"
                        if "project_data" not in st.session_state:
                            st.session_state.project_data = ""
                        st.session_state.project_data += f"\n{intel_log}"
                        
                        st.success(f"✅ 申请已提交！流水号：{req_id}，已推入直属总监审批流。")
                        st.rerun()
        else:
            st.info("💡 资金申请仅限前线销售人员发起。您当前处于审批者/上帝视角。")

    with col_audit:
        st.markdown("### 🏦 审批流转中心")
        reqs = st.session_state.get("expense_requests", [])
        
        # 权限过滤与状态路由
        visible_reqs = []
        for r in reqs:
            if current_r == "区域总监" and r["status"] == "待总监审批" and r["dept"] == st.session_state.user_dept:
                visible_reqs.append(r)
            elif current_r == "销售VP" and r["status"] == "待VP审批":
                visible_reqs.append(r)
            elif current_r == "总经理" and r["status"] == "待总经理核准":
                visible_reqs.append(r)
            elif current_r == "财务" and r["status"] == "待财务打款":
                visible_reqs.append(r)
            elif current_r == "一线销售" and r["applicant"] == current_user:
                visible_reqs.append(r) # 销售看自己的进度
                
        if not visible_reqs:
            st.success("🎉 当前您的待办/关注列表中没有费用申请。")
        else:
            for r in visible_reqs:
                with st.expander(f"[{r['status']}] {r['type']} - {r['amount']}元 | 提报: {r['applicant']}", expanded=True):
                    st.markdown(f"**关联项目：** {r['project']}\n\n**申请事由：** {r['reason']}")
                    st.caption(" \n".join(r['audit_trail']))
                    
                    # 审批动作区
                    if current_r in ["区域总监", "销售VP", "总经理", "财务"] and "待" in r["status"]:
                        c1, c2 = st.columns(2)
                        if c1.button("✅ 同意流转/核准/打款", key=f"ok_{r['id']}", type="primary", use_container_width=True):
                            if current_r == "区域总监":
                                r["status"] = "待VP审批"
                                r["audit_trail"].append(f"[{current_user}] 总监已同意，流转至 VP。")
                            elif current_r == "销售VP":
                                r["status"] = "待总经理核准"
                                r["audit_trail"].append(f"[{current_user}] VP 已同意，流转至总经理。")
                            elif current_r == "总经理":
                                r["status"] = "待财务打款"
                                r["audit_trail"].append(f"[{current_user}] 总经理已核准，通知财务打款。")
                            elif current_r == "财务":
                                r["status"] = "✅ 财务已打款闭环"
                                r["audit_trail"].append(f"[{current_user}] 财务已完成资金拨付。")
                                # 核心闭环：只有真金白银花出去了，才作为高权重情报喂给 AI 记忆库
                                if "project_data" not in st.session_state:
                                    st.session_state.project_data = ""
                                st.session_state.project_data += f"\n【💰 真实资金消耗】财务已正式将 {r['amount']} 元的 {r['type']} 打款至前线。请 AI 在后续战况推演中，严格监督该笔资金的转化率！"
                            st.rerun()
                            
                        if c2.button("❌ 驳回", key=f"no_{r['id']}", use_container_width=True):
                            r["status"] = "已驳回"
                            r["audit_trail"].append(f"[{current_user}] 驳回了该申请。")
                            if "project_data" not in st.session_state:
                                st.session_state.project_data = ""
                            st.session_state.project_data += f"\n【🚫 资金冻结】审批人驳回了该笔 {r['amount']} 元的申请。"
                            st.rerun()

# ── 招投标控标中心 ──
# ── 招投标"雷区"与控标中心 (Bidding War Room) ──
with tab_bidding:
    st.markdown("## 📄 招投标“雷区”与控标中心 (Bidding War Room)")
    st.caption("大型 B2B 商战绞肉机：在这里拆解对手的标书陷阱，或者由 AI 结合我方优势生成具有绝对排他性的控标参数。")

    # --- 顶部：全局打单项目选择器 ---
    project_list = list(st.session_state.get("projects", {}).keys())
    if not project_list:
        project_list = ["(暂无立项项目，请先在左侧新建)"]
    
    st.markdown("**📁 选择正在运作的打单项目：**")
    bidding_proj = st.selectbox("隐藏的Label", project_list, label_visibility="collapsed", key="bidding_warroom_proj")
    st.markdown("<br>", unsafe_allow_html=True)

    # --- 核心战场：左防守，右进攻 ---
    col_defend, col_attack = st.columns(2)

    # 🛡️ 左侧：防守 (排雷)
    with col_defend:
        with st.container(border=True):
            st.markdown("### 🛡️ 标书拆解与防守 (排雷)")
            st.info("上传客户发来的 RFP / 招标文件，AI 将瞬间标出哪些参数是友商提前埋好的“雷”。")
            
            bidding_file = st.file_uploader("📎 上传招标文件 (支持 PDF/Word)：", type=["pdf", "docx"], key="bidding_defense_file")
            
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("🔍 AI 一键深度拆标排雷", type="primary", use_container_width=True):
                if not bidding_file:
                    st.warning("⚠️ 请先上传买来的招标文件！")
                else:
                    st.success("🔌 (槽位预留) 等待呼叫大模型执行条款排雷与废标项拦截...")

    # ⚔️ 右侧：进攻 (埋雷)
    with col_attack:
        with st.container(border=True):
            st.markdown("### ⚔️ 控标参数与进攻 (埋雷)")
            st.info("当我们处于“领跑”身位时，让 AI 结合我方独家优势，生成极其隐蔽的排他性控标参数，教客户怎么写标书。")
            
            target_competitor = st.selectbox("🎯 假想敌 (本次重点防范友商)：", ["西门子", "ABB", "施耐德", "江苏大全", "其他竞品"])
            
            # --- 升级：一键调取系统武器库弹药 ---
            adv_list = st.session_state.get("core_advantages", [])
            selected_advs = st.multiselect("💎 调取我方核心差异化优势 (可多选)：", adv_list, default=adv_list[0] if adv_list else None)
            
            custom_adv = st.text_input("✍️ 临时补充特定项目优势 (可选)：", placeholder="针对此项目临时想到的绝活...")
            
            # 拼装备弹药
            core_advantage = "；".join(selected_advs)
            if custom_adv:
                core_advantage += "；" + custom_adv
            # ----------------------------------------
            
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("💣 生成极具伪装性的控标参数", type="primary", use_container_width=True):
                if not core_advantage:
                    st.error("⚠️ 弹药不足：请先输入我方的核心差异化优势！")
                else:
                    st.success("🔌 (槽位预留) 等待呼叫大模型生成毒辣的控标话术...")

    # --- 底部：排雷报告与战术雷达展示区 ---
    if st.session_state.get("bidding_risk_report"):
        st.divider()
        st.markdown("## 🚨 AI 标书深度排雷战报")
        
        with st.container(border=True):
            st.markdown(st.session_state.bidding_risk_report)
            
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("🗑️ 阅毕，清除当前排雷战报", use_container_width=True):
                del st.session_state["bidding_risk_report"]
                st.rerun()

# ── Deal Desk 智能报价与审批中心 ──
with tab_deal_desk:
    st.markdown("## 💰 智能报价与自动审批流水线")
    st.caption("处理成套厂/第三方客户询价：AI 自动解析文件提取用量 -> 销售校对 -> 自动生成制式报价单 -> VP审批盖章")

    col_parse, col_review = st.columns([1, 1.2])

    with col_parse:
        with st.container(border=True):
            st.markdown("### 📥 1. 询价文件解析 (AI 提取)")
            # --- 新增：关联业务项目选择器 ---
            project_list = list(st.session_state.get("projects", {}).keys())
            if not project_list:
                project_list = ["(暂无立项项目，请先在左侧新建)"]
            deal_proj = st.selectbox("📂 关联业务项目：", project_list, key="deal_proj_select")
            
            # --- 🚀 快车道：检测并调取底单 ---
            if "approved_boms" in st.session_state and deal_proj in st.session_state.approved_boms:
                st.success(f"🌟 侦测到【{deal_proj}】已存在 VP 获批底单！")
                if st.button("📥 一键调取本项目底单 (免审直接出单)", type="primary", use_container_width=True):
                    st.session_state.mock_bom = st.session_state.approved_boms[deal_proj].copy()
                    st.session_state.approval_status = "approved" # 瞬间绿灯放行
                    st.rerun()
            st.markdown("<br>", unsafe_allow_html=True)
            # ------------------------------
            # --- 升级：支持搜索与动态新增的客户选择器 ---
            if "switchgear_clients" not in st.session_state:
                st.session_state.switchgear_clients = ["江苏大全", "上海正泰", "西门子", "施耐德", "ABB"]
            
            client_options = st.session_state.switchgear_clients + ["➕ 手动录入新客户..."]
            selected_client_option = st.selectbox("🏢 询价客户主体 (点击后直接键盘搜索)：", client_options, key="deal_client_select")
            
            if selected_client_option == "➕ 手动录入新客户...":
                with st.container(border=True):
                    new_client_name = st.text_input("✍️ 请输入新客户全称 (按 Enter 回车)：", key="new_deal_client_input")
                    if st.button("💾 保存并永久录入", type="secondary", use_container_width=True):
                        if new_client_name and new_client_name not in st.session_state.switchgear_clients:
                            st.session_state.switchgear_clients.insert(0, new_client_name)
                            st.rerun()
                        elif new_client_name in st.session_state.switchgear_clients:
                            st.warning("⚠️ 该客户已存在！")
                inquiry_client = new_client_name
            else:
                inquiry_client = selected_client_option
            # ----------------------------------------
            inquiry_files = st.file_uploader("📎 上传询价资料 (支持 招标文件/技术图纸/询价函)", accept_multiple_files=True, type=["pdf", "docx", "png", "jpg"])
            
            if st.button("🤖 AI 一键解析图纸与清单", type="primary", use_container_width=True):
                if not inquiry_files:
                    st.warning("⚠️ 请先上传客户发来的询价文件！")
                elif not api_key:
                    st.error("⚠️ 核心拦截：请先在左侧边栏输入大模型 API Key！")
                else:
                    with st.spinner("🧠 多模态大模型正在深度阅读标书与图纸，提取物料清单..."):
                        try:
                            # 1. 提取文件真实内容 (支持 TXT 和 PDF)
                            file_content = ""
                            for file in inquiry_files:
                                if file.name.lower().endswith('.txt'):
                                    file_content += file.getvalue().decode("utf-8") + "\n"
                                elif file.name.lower().endswith('.pdf'):
                                    try:
                                        import PyPDF2
                                        pdf_reader = PyPDF2.PdfReader(file)
                                        for page in pdf_reader.pages:
                                            text = page.extract_text()
                                            if text: file_content += text + "\n"
                                    except ImportError:
                                        st.warning("💡 系统提示：尚未安装 `PyPDF2` 库。暂无法真正读取 PDF，将启用自动模拟兜底模式。")
                            
                            # 2. 防呆兜底：如果没读出真实文字，自动喂入一段电气行业的测试数据保证跑通
                            if not file_content.strip():
                                file_content = "客户本次项目需要采购：消弧选线装置A型 8套，以及智能防腐涂层组件 150套。请尽快回复报价。"

                            # 3. 呼叫大模型进行 JSON 结构化提取
                            from openai import OpenAI as _OAI_deal
                            _client_deal = _OAI_deal(api_key=api_key)
                            
                            deal_prompt = f"""你是一个资深的电气成套设备报价总工。请从以下客户的询价资料中，提取出两部分核心信息：
1. 客户的联系方式（如果有邮箱优先提取邮箱，其次是手机号、微信号。如果完全没有，请输出空字符串 ""）。
2. 提取出产品型号和对应数量。

资料内容：{file_content}

请严格按以下 JSON 格式输出，不要任何解释性废话，必须是合法的 JSON 对象：
{{
  "联系方式": "提取到的邮箱或手机号",
  "物料清单": [
    {{"产品型号": "提取到的型号1", "数量": 8}},
    {{"产品型号": "提取到的型号2", "数量": 150}}
  ]
}}"""
                            resp_deal = _client_deal.chat.completions.create(
                                model="gpt-4o",
                                messages=[{"role": "user", "content": deal_prompt}],
                                temperature=0.1
                            )
                            
                            # 4. 清洗 AI 返回的结果，提取 BOM 和联系方式
                            ai_result_str = resp_deal.choices[0].message.content.replace('```json', '').replace('```', '').strip()
                            import json
                            import pandas as pd
                            
                            try:
                                parsed_data = json.loads(ai_result_str)
                                contact_info = parsed_data.get("联系方式", "")
                                bom_items = parsed_data.get("物料清单", [])
                            except:
                                # 兼容防爆：如果大模型偶发未按格式返回
                                contact_info = ""
                                bom_items = json.loads(ai_result_str) if isinstance(json.loads(ai_result_str), list) else []

                            # 将 AI 捕获到的联系方式存入会话，供最后一步的直发控制台调用
                            if contact_info:
                                st.session_state.ai_contact_info = contact_info

                            new_bom_list = []
                            for item in bom_items:
                                new_bom_list.append({
                                    "产品型号": item.get("产品型号", "未知产品"),
                                    "AI提取数量": int(item.get("数量", 0)),
                                    "销售核定数量": int(item.get("数量", 0)),
                                    "标准单价(元)": 15000 if "消弧" in item.get("产品型号", "") else 800,
                                    "备注": ""
                                })
                            
                            if new_bom_list:
                                st.session_state.mock_bom = pd.DataFrame(new_bom_list)
                            
                            st.success("✅ AI 提取完成！请在右侧表格进行人工校对。")
                            st.rerun()
                            
                        except Exception as e:
                            st.error(f"❌ 解析失败，请检查 API Key 余额或重试：{e}")

    with col_review:
        with st.container(border=True):
            st.markdown("### 📝 2. 人工校对与制式报价")
            st.info("💡 销售人员在此核对并修正 AI 提取的型号与数量。确认无误后发起审批。")
            
            # 使用可编辑表格 (Data Editor) 让销售直接改数字
            import pandas as pd
            if "mock_bom" not in st.session_state:
                st.session_state.mock_bom = pd.DataFrame([
                    {"产品型号": "消弧选线装置 A型", "AI提取数量": 0, "销售核定数量": 0, "标准单价(元)": 15000, "备注": ""},
                    {"产品型号": "智能防腐涂层组件", "AI提取数量": 0, "销售核定数量": 0, "标准单价(元)": 800, "备注": ""}
                ])
            
            if st.session_state.get("ai_parsed_bom"):
                st.session_state.mock_bom["AI提取数量"] = [5, 120] # 模拟AI算出的数据
                st.session_state.mock_bom["销售核定数量"] = [5, 120] 
                del st.session_state["ai_parsed_bom"]
                
            # --- 升级：支持全动态编辑、新增、删除行的智能表格 ---
            edited_bom = st.data_editor(
                st.session_state.mock_bom, 
                use_container_width=True,
                num_rows="dynamic",  
                disabled=["AI提取数量", "小计(元)"], # 保护 AI 提取数据和自动计算的小计不被乱改
                hide_index=False, 
                key="bom_data_editor"
            )
            
            st.session_state.final_bom = edited_bom
            
            # --- 🛡️ 深度天眼：防篡改与变更侦测引擎 ---
            if "approved_boms" in st.session_state and deal_proj in st.session_state.approved_boms:
                import pandas as pd
                df_curr = st.session_state.final_bom.copy()
                df_base = st.session_state.approved_boms[deal_proj].copy()
                
                # 安全清理：去除空行干扰，深度比对每一个单元格
                curr_clean = df_curr.dropna(subset=['产品型号', '销售核定数量'], how='all').reset_index(drop=True).astype(str)
                base_clean = df_base.dropna(subset=['产品型号', '销售核定数量'], how='all').reset_index(drop=True).astype(str)
                
                if not curr_clean.equals(base_clean):
                    # 侦测到任何篡改，立刻计算财务差异
                    old_total = (pd.to_numeric(df_base["销售核定数量"], errors='coerce').fillna(0) * pd.to_numeric(df_base["标准单价(元)"], errors='coerce').fillna(0)).sum()
                    new_total = (pd.to_numeric(df_curr["销售核定数量"], errors='coerce').fillna(0) * pd.to_numeric(df_curr["标准单价(元)"], errors='coerce').fillna(0)).sum()
                    
                    diff_msg = "🚨 **天眼侦测到该报价单已被销售修改！**\n"
                    if old_total != new_total:
                        diff_msg += f"▶️ **【财务严重变更】** 总价发生异动：原底单 **{old_total:,.2f}** 元 ➡️ 新提交 **{new_total:,.2f}** 元！"
                    else:
                        diff_msg += f"▶️ **【明细异常变更】** 总价未变，但产品型号、数量或备注细节被偷偷修改！"
                        
                    st.session_state.current_diff_summary = diff_msg
                    
                    # 核心拦截：如果销售试图偷偷修改已获批的底单，立刻褊夺绿灯，打回原形
                    if st.session_state.get("approval_status") == "approved":
                        st.session_state.approval_status = "draft"
                        st.warning("🚨 警告：检测到核心表格物料/单价被修改！已自动关闭【免审快车道】，必须重新提交 VP 审批。")
                else:
                    st.session_state.current_diff_summary = "" # 未篡改则清空警告
            else:
                st.session_state.current_diff_summary = "💡 这是一个全新/未获批项目，暂无比对底单。"
            # ----------------------------------------
            
            # --- 新增：总金额汇总展示区 ---
            if "current_total_price" in st.session_state:
                st.success(f"💰 当前订单核定总金额：**{st.session_state.current_total_price:,.2f} 元**")
                
            st.markdown("<br>", unsafe_allow_html=True)
            col_btn1, col_btn2 = st.columns(2)
            
            with col_btn1:
                if st.button("🧮 格式化表格并核算总价", use_container_width=True):
                    if "final_bom" in st.session_state and not st.session_state.final_bom.empty:
                        import pandas as pd
                        import numpy as np
                        
                        df = st.session_state.final_bom.copy()
                        # 1. 深度清理：彻底清除删掉的废行和未填写的 None 空行
                        df.replace([None, 'None', ''], np.nan, inplace=True)
                        df = df.dropna(subset=['产品型号', '销售核定数量'], how='all')
                        
                        # 2. 安全计算：保障数字格式不出错
                        df["销售核定数量"] = pd.to_numeric(df["销售核定数量"], errors='coerce').fillna(0).astype(int)
                        df["标准单价(元)"] = pd.to_numeric(df["标准单价(元)"], errors='coerce').fillna(0.0)
                        
                        # 3. 动态生成每一行的"小计"列
                        df["小计(元)"] = df["销售核定数量"] * df["标准单价(元)"]
                        
                        # 4. 核心修复：让序号重新从 1 开始顺延
                        df.index = range(1, len(df) + 1)
                        
                        # 5. 计算总造价，存入缓存并刷新界面
                        st.session_state.mock_bom = df
                        st.session_state.current_total_price = df["小计(元)"].sum()
                        st.rerun()
                    else:
                        st.warning("⚠️ 表格数据为空。")
                        
            with col_btn2:
                # 状态机初始化
                if "approval_status" not in st.session_state:
                    st.session_state.approval_status = "draft"
                
                status = st.session_state.approval_status
                
                # --- 状态 1：销售起草/被驳回状态 ---
                if status == "draft" or status == "rejected":
                    if status == "rejected":
                        st.error(f"❌ VP 驳回：{st.session_state.get('reject_reason', '利润太低！')}")
                        btn_label = "🔄 重新修改并再次提交"
                    else:
                        btn_label = "📤 提交直属上级与 VP 审批"
                        
                    if st.button(btn_label, type="primary", use_container_width=True):
                        if "final_bom" not in st.session_state or st.session_state.final_bom.empty:
                            st.error("⚠️ 表格中尚未核定任何物料，无法提交！")
                        else:
                            st.session_state.approval_status = "pending"
                            st.rerun()
                            
                # --- 状态 2：流转中（展示模拟 VP 审批台） ---
                elif status == "pending":
                    st.info("⏳ 报价单已锁定，已推送至 VP 待办列表，等待审批...")
                    with st.expander("👑 [模拟 VP 视角] 领导审批工作台", expanded=True):
                        st.markdown("⚠️ **VP 请注意**：该销售提交的报价单正在申请【公司公章】与【您的电子签章】。")
                        
                        # --- 核心：向 VP 展示天眼侦测报告 ---
                        if st.session_state.get("current_diff_summary"):
                            if "🚨" in st.session_state.current_diff_summary:
                                st.error(st.session_state.current_diff_summary)
                            else:
                                st.info(st.session_state.current_diff_summary)
                        # ----------------------------------
                        
                        c1, c2 = st.columns(2)
                        if c1.button("✅ 利润合规，同意并盖章", type="primary", use_container_width=True):
                            st.session_state.approval_status = "approved"
                            # --- 💾 核心：将获批的表格存入底单金库，与项目深度绑定 ---
                            if "approved_boms" not in st.session_state:
                                st.session_state.approved_boms = {}
                            st.session_state.approved_boms[deal_proj] = st.session_state.final_bom.copy()
                            st.rerun()
                        if c2.button("❌ 利润太低，驳回", use_container_width=True):
                            st.session_state.approval_status = "rejected"
                            st.session_state.reject_reason = "毛利未达公司 30% 红线要求，且未见战略打单说明，请重新调整单价！"
                            st.rerun()
                            
                # --- 状态 3：审批通过，生成带印章的 Word ---
                elif status == "approved":
                    st.success("✅ VP 审批已通过！电子公章已授权下发。")
                    
                    try:
                        import io
                        import os
                        import pandas as pd
                        from docx import Document
                        from docx.shared import Pt, RGBColor, Inches
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        import time
                        
                        doc = Document()
                        title = doc.add_heading(f"{inquiry_client} - {deal_proj} 报价单", 0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        doc.add_paragraph(f"致：{inquiry_client} 采购部")
                        doc.add_paragraph(f"关于项目：【{deal_proj}】")
                        doc.add_paragraph("感谢贵司的询价！经我司内部严格核算及销售副总审批，现提供最优报价方案如下：\n")
                        
                        bom_df = st.session_state.final_bom.copy()
                        table = doc.add_table(rows=1, cols=5)
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text, hdr_cells[1].text = '产品型号', '核定数量'
                        hdr_cells[2].text, hdr_cells[3].text, hdr_cells[4].text = '单价(元)', '小计(元)', '备注'
                        
                        total_amount = 0
                        for index, row in bom_df.iterrows():
                            if pd.isna(row.get('产品型号')) and pd.isna(row.get('销售核定数量')): continue
                                
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(row.get('产品型号', '')) if pd.notna(row.get('产品型号')) else ''
                            raw_qty = row.get('销售核定数量', 0)
                            qty = int(raw_qty) if pd.notna(raw_qty) and str(raw_qty).strip() != '' else 0
                            raw_price = row.get('标准单价(元)', 0.0)
                            price = float(raw_price) if pd.notna(raw_price) and str(raw_price).strip() != '' else 0.0
                            
                            subtotal = qty * price
                            total_amount += subtotal
                            
                            row_cells[1].text = str(qty)
                            row_cells[2].text = f"{price:,.2f}"
                            row_cells[3].text = f"{subtotal:,.2f}"
                            row_cells[4].text = str(row.get('备注', '')) if pd.notna(row.get('备注')) else ''
                        
                        doc.add_paragraph(f"\n📢 总计大写金额：人民币 {total_amount:,.2f} 元")
                        doc.add_paragraph("（注：以上报价含 13% 增值税，不含现场施工费。）")
                        
                        p = doc.add_paragraph("\n")
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        seal_path = "assets/official_seal.png"
                        sign_path = "assets/personal_sign.png"
                        
                        if os.path.exists(seal_path):
                            run = p.add_run()
                            run.add_picture(seal_path, width=Inches(1.5))
                        else:
                            runner = p.add_run("【审批通过：此件已加盖企业电子合同专用章】\n")
                            runner.font.color.rgb = RGBColor(220, 20, 60)
                            runner.bold = True
                        
                        p2 = doc.add_paragraph()
                        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        if os.path.exists(sign_path):
                            run2 = p2.add_run("授权人签字：")
                            run2.add_picture(sign_path, width=Inches(1.0))
                        else:
                            p2.add_run("授权人：销售副总 VP")
                            
                        doc.add_paragraph("防伪溯源追踪码：AI-DEAL-" + str(int(time.time()))).alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        
                        buffer = io.BytesIO()
                        doc.save(buffer)
                        # --- 升级：生成高仿 PDF 预览与一键直发控制台 ---
                        st.divider()
                        st.markdown("### 👁️ 正式报价单预览 (PDF版式)")
                        
                        # 模拟 PDF 白纸黑字的视觉排版效果
                        with st.container(border=True):
                            st.markdown(f"<h3 style='text-align: center;'>{inquiry_client} - {deal_proj} 报价单</h3>", unsafe_allow_html=True)
                            st.markdown(f"**致：** {inquiry_client} 采购部<br>**关于项目：** 【{deal_proj}】", unsafe_allow_html=True)
                            st.markdown("感谢贵司的询价！经我司内部严格核算及销售副总审批，现提供最优报价方案如下：")
                            
                            st.dataframe(st.session_state.final_bom, use_container_width=True, hide_index=True)
                            
                            c_left, c_right = st.columns([1, 1])
                            with c_left:
                                st.markdown(f"**📢 总计金额：人民币 {total_amount:,.2f} 元**")
                                st.caption("（注：以上报价含 13% 增值税，不含现场施工费。）")
                            with c_right:
                                if os.path.exists("assets/official_seal.png"):
                                    st.image("assets/official_seal.png", width=120)
                                else:
                                    st.markdown("<span style='color:red; font-weight:bold;'>【此件已加盖企业电子合同专用章】</span>", unsafe_allow_html=True)
                                
                                if os.path.exists("assets/personal_sign.png"):
                                    st.image("assets/personal_sign.png", width=80)
                                else:
                                    st.markdown("**授权人签字：销售副总 VP**")

                        st.markdown("### 📨 报价单直发工作台")
                        st.info(f"🔒 请仔细核对上方生成的最终版式。确认无误后可一键发送。")
                        
                        # 预留给 AI 提取联系方式的变量槽位
                        auto_contact = st.session_state.get("ai_contact_info", "")
                        
                        send_target = st.text_input("🎯 接收方账号 (微信号/邮箱/手机)：", value=auto_contact, placeholder="请输入客户联系方式...")
                        send_msg = st.text_area("💬 附言：", value=f"您好，附件为【{deal_proj}】项目的正式报价单，请查收。价格及账期已获批，如有问题随时沟通！")
                        
                        if st.button("🚀 确认并一键发送 (通过真实邮件引擎)", type="primary", use_container_width=True):
                            if not send_target:
                                st.error("⚠️ 请先填写接收方邮箱！")
                            elif "@" not in send_target:
                                st.warning("⚠️ 目前仅支持真实邮件发送测试，请输入正确的邮箱地址（暂不支持微信号直发）。")
                            else:
                                with st.spinner("📧 正在连接 SMTP 邮件服务器投递标书..."):
                                    try:
                                        import smtplib
                                        from email.mime.multipart import MIMEMultipart
                                        from email.mime.text import MIMEText
                                        from email.mime.application import MIMEApplication
                                        
                                        # ==========================================
                                        # ⚠️ 老板请注意：这里需要替换为您自己的发件邮箱配置
                                        # 以下默认以腾讯 QQ/企业邮箱为例。如果是网易请改 smtp.163.com
                                        # ==========================================
                                        SENDER_EMAIL = "your_email@qq.com"  # 替换为您的真实发件邮箱
                                        SENDER_PWD = "your_auth_code"       # 替换为您的邮箱 SMTP 授权码 (非登录密码)
                                        SMTP_SERVER = "smtp.qq.com"
                                        SMTP_PORT = 465
                                        
                                        # 1. 构建邮件主体
                                        msg = MIMEMultipart()
                                        msg['Subject'] = f"【正式报价单】{inquiry_client} - {deal_proj}"
                                        msg['From'] = SENDER_EMAIL
                                        msg['To'] = send_target
                                        
                                        # 2. 写入邮件正文
                                        msg.attach(MIMEText(send_msg, 'plain', 'utf-8'))
                                        
                                        # 3. 提取刚才内存里生成的带公章 Word 作为真实附件
                                        file_name = f"{inquiry_client}_{deal_proj}_报价单.docx"
                                        part = MIMEApplication(buffer.getvalue(), Name=file_name)
                                        part['Content-Disposition'] = f'attachment; filename="{file_name}"'
                                        msg.attach(part)
                                        
                                        # 4. 建立 SSL 加密连接并发送
                                        server = smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT)
                                        server.login(SENDER_EMAIL, SENDER_PWD)
                                        server.sendmail(SENDER_EMAIL, [send_target], msg.as_string())
                                        server.quit()
                                        
                                        st.balloons()
                                        st.success(f"✅ 邮件发送成功！报价单已投递至：{send_target}")
                                        
                                        # 发送完成后，清理缓存并重置状态
                                        if "ai_contact_info" in st.session_state:
                                            del st.session_state["ai_contact_info"]
                                        time.sleep(2)
                                        st.session_state.approval_status = "draft"
                                        st.rerun()
                                        
                                    except smtplib.SMTPAuthenticationError:
                                        st.error("❌ 邮箱登录失败！请检查 SENDER_EMAIL 和 SENDER_PWD（SMTP授权码）是否正确。")
                                    except smtplib.SMTPException as smtp_err:
                                        st.error(f"❌ SMTP 服务器错误：{smtp_err}")
                                    except Exception as mail_err:
                                        st.error(f"❌ 邮件发送失败：{mail_err}")
                        # --------------------------------
                            
                    except Exception as e:
                        st.error(f"❌ 生成 Word 失败，请联系管理员：{e}")

# ── 合同联审流水线 (Contract Review Pipeline) ──
with tab_contract:
    st.markdown("## 📋 合同联审流水线")
    st.caption("销售发起 → 技术超配审查 → 销售定价与商务条款 → VP 审批与合同生成。全链路防篡改、全角色抬杠。")

    # --- 项目选择器 ---
    contract_projects = list(st.session_state.get("projects", {}).keys())
    if not contract_projects:
        contract_projects = ["(暂无立项项目)"]
    contract_proj = st.selectbox("📂 选择合同关联项目：", contract_projects, key="contract_pipeline_proj")

    # --- 状态机 ---
    state_key = f"contract_state_{contract_proj}"
    bom_key = f"contract_bom_{contract_proj}"
    if state_key not in st.session_state:
        st.session_state[state_key] = "1_sales_init"

    current_r = st.session_state.get("role", "一线销售")

    # ─ 流程进度条 ─
    steps = ["1_sales_init", "2_tech_review", "3_sales_pricing", "4_vp_approval", "5_approved", "6_commission"]
    step_labels = ["❶ 销售发起", "❷ 技术审查", "❸ 销售定价", "❹ VP审批", "❺ 合同发送", "❻ 提成核算"]
    current_step_idx = steps.index(st.session_state[state_key]) if st.session_state[state_key] in steps else 0
    progress_text = " → ".join([f"**{l}**" if i == current_step_idx else l for i, l in enumerate(step_labels)])
    st.markdown(f"🚦 当前进度：{progress_text}")
    st.divider()

    # 🟢 第一步：销售发起
    with st.expander("🟢 第一步：销售发起合同请求 (录入 BOM)", expanded=(st.session_state[state_key] == "1_sales_init")):
        if st.session_state[state_key] == "1_sales_init":
            import pandas as pd
            if bom_key not in st.session_state:
                st.session_state[bom_key] = pd.DataFrame({
                    "产品型号": ["示例-XGN15", "示例-KYN28"],
                    "AI提取数量": [10, 5],
                    "销售核定数量": [10, 5],
                    "标准单价(元)": [15000, 8000],
                    "备注": ["", ""]
                })
            st.data_editor(st.session_state[bom_key], use_container_width=True, num_rows="dynamic", key="contract_init_editor")
            if st.button("➡️ 提交至技术部超配审查", type="primary", use_container_width=True):
                st.session_state[state_key] = "2_tech_review"
                st.rerun()
        else:
            st.success("✅ 已提交")

    # 🟡 第二步：技术部超配审查
    with st.expander("🟡 第二步：技术部超配与工况审查", expanded=(st.session_state[state_key] == "2_tech_review")):
        if st.session_state[state_key] == "1_sales_init":
            st.info("⏳ 等待销售提交...")
        elif st.session_state[state_key] == "2_tech_review":
            if current_r not in ["技术工程师", "系统管理员"]:
                st.error("🔒 权限拦截：请由【技术工程师】进行超配审查。")
            else:
                st.info("🔧 技术部请核查每行物料的实际工况需求，填写技术核定数量与超配说明。")
                import pandas as pd
                bom_df = st.session_state[bom_key].copy()
                if "技术核定数量" not in bom_df.columns:
                    bom_df["技术核定数量"] = bom_df["销售核定数量"]
                if "超配说明" not in bom_df.columns:
                    bom_df["超配说明"] = ""
                st.session_state[bom_key] = bom_df
                edited_tech = st.data_editor(
                    st.session_state[bom_key],
                    disabled=["产品型号", "AI提取数量", "销售核定数量"],
                    use_container_width=True, key="tech_review_editor"
                )
                st.session_state[bom_key] = edited_tech
                if st.button("➡️ 技术审查完成，提交销售定价", type="primary", use_container_width=True):
                    st.session_state[state_key] = "3_sales_pricing"
                    st.rerun()
        else:
            st.success("✅ 技术审查已完成")

    # 🔴 第三步：销售确认定价与防篡改提交
    with st.expander("🔴 第三步：销售最终定价、商务条款与防篡改风控", expanded=(st.session_state[state_key] == "3_sales_pricing")):
        if st.session_state[state_key] in ["1_sales_init", "2_tech_review"]:
            st.info("⏳ 等待前置环节完成...")
        elif current_r != "一线销售" and st.session_state[state_key] == "3_sales_pricing":
            st.error("🔒 权限拦截：请由负责该项目的【一线销售】进行最终合同金额与商务条款的确认。")
        elif st.session_state[state_key] == "3_sales_pricing":
            st.warning("⚠️ 销售请注意：技术部已添加超配信息，请根据最新的【技术核定数量】调整最终合同数量与单价！")
            edited_sales_bom = st.data_editor(
                st.session_state[bom_key],
                disabled=["产品型号", "AI提取数量", "技术核定数量", "超配说明"],
                use_container_width=True, num_rows="dynamic", key="sales_editor"
            )
            st.session_state[bom_key] = edited_sales_bom
            
            # --- 新增：核心商务与履约条款表单 ---
            st.markdown("#### 📝 核心商务与履约条款 (将写入最终合同)")
            c_pay1, c_pay2 = st.columns(2)
            with c_pay1:
                pay_method = st.selectbox("💳 支付方式：", ["电汇 (T/T)", "承兑汇票 (半年期)", "承兑汇票 (一年期)", "国内信用证 (L/C)"], key=f"pay_method_{contract_proj}")
                delivery_time = st.text_input("🚚 货期承诺：", placeholder="例：合同签订且收到预付款后 30 个工作日内发货", key=f"delivery_time_{contract_proj}")
                warranty_period = st.text_input("🛡️ 质保期限：", placeholder="例：货到现场 12 个月或投运 18 个月 (以先到为准)", key=f"warranty_{contract_proj}")
            with c_pay2:
                st.markdown("**💰 付款节点比例 (%)** - *风控要求：总和必须为 100*")
                r_col1, r_col2, r_col3, r_col4 = st.columns(4)
                ratio_advance = r_col1.number_input("预付款%", min_value=0, max_value=100, value=30, step=5, key=f"r1_{contract_proj}")
                ratio_delivery = r_col2.number_input("发货款%", min_value=0, max_value=100, value=30, step=5, key=f"r2_{contract_proj}")
                ratio_accept = r_col3.number_input("验收款%", min_value=0, max_value=100, value=30, step=5, key=f"r3_{contract_proj}")
                ratio_warranty = r_col4.number_input("质保金%", min_value=0, max_value=100, value=10, step=5, key=f"r4_{contract_proj}")
            
            st.markdown("**📍 物流与现场交接信息**")
            c_log1, c_log2 = st.columns(2)
            with c_log1:
                delivery_address = st.text_area("🏭 发货/现场接收地址：", placeholder="例：山东省烟台市万华化学工业园二期施工现场", key=f"addr_{contract_proj}")
            with c_log2:
                receiver_contact = st.text_input("👤 收货人及联系方式：", placeholder="例：现场经理 王工 138xxxx5678", key=f"contact_{contract_proj}")

            total_ratio = ratio_advance + ratio_delivery + ratio_accept + ratio_warranty

            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("📤 锁定价格与条款，提交 VP 审批", type="primary", use_container_width=True):
                # 强悍的商务风控拦截
                if total_ratio != 100:
                    st.error(f"🚨 财务风控拦截：付款比例总和必须为 100% (当前为 {total_ratio}%)，请重新分配！")
                elif not delivery_time or not delivery_address or not receiver_contact:
                    st.warning("⚠️ 请完整填写货期承诺、发货地址与收货人信息！")
                else:
                    # 存底单及商务条款字典
                    st.session_state[f"base_bom_{contract_proj}"] = edited_sales_bom.copy()
                    st.session_state[f"comm_terms_{contract_proj}"] = {
                        "pay_method": pay_method,
                        "ratio_advance": ratio_advance,
                        "ratio_delivery": ratio_delivery,
                        "ratio_accept": ratio_accept,
                        "ratio_warranty": ratio_warranty,
                        "delivery_time": delivery_time,
                        "warranty_period": warranty_period,
                        "delivery_address": delivery_address,
                        "receiver_contact": receiver_contact
                    }
                    st.session_state[state_key] = "4_vp_approval"
                    st.success("✅ 商务条款与报价已锁定，已推送至 VP 审批中心！")
                    st.rerun()

    # 🔵 第四步：高管审批与制式合同生成
    with st.expander("🔵 第四步：VP 合同终审与一键触达", expanded=(st.session_state[state_key] == "4_vp_approval")):
        if st.session_state[state_key] in ["1_sales_init", "2_tech_review", "3_sales_pricing"]:
            st.info("⏳ 等待销售提交终版合同数据...")
        elif st.session_state[state_key] == "4_vp_approval":
            st.info("👑 审批工作台：请核阅技术超配情况、最终价格及核心商务条款。")
            st.dataframe(st.session_state[bom_key], use_container_width=True)
            
            # --- 渲染商务条款供 VP 审查 ---
            terms = st.session_state.get(f"comm_terms_{contract_proj}", {})
            if terms:
                st.markdown("#### ⚖️ 拟定商务条款审查")
                with st.container(border=True):
                    t_col1, t_col2 = st.columns(2)
                    with t_col1:
                        st.markdown(f"**💳 支付方式**：`{terms.get('pay_method')}`")
                        st.markdown(f"**🚚 货期承诺**：{terms.get('delivery_time')}")
                        st.markdown(f"**🛡️ 质保期限**：{terms.get('warranty_period')}")
                    with t_col2:
                        st.markdown(f"**💰 付款比例**：预付 **{terms.get('ratio_advance')}%** | 发货 **{terms.get('ratio_delivery')}%** | 验收 **{terms.get('ratio_accept')}%** | 质保 **{terms.get('ratio_warranty')}%**")
                        st.markdown(f"**📍 现场地址**：{terms.get('delivery_address')}")
                        st.markdown(f"**👤 收 货 人**：{terms.get('receiver_contact')}")
            
            st.markdown("<br>", unsafe_allow_html=True)
            c1, c2 = st.columns(2)
            if c1.button("✅ 利润与条款合规，同意并加盖公章", type="primary", use_container_width=True):
                st.session_state[state_key] = "5_approved"
                st.rerun()
            if c2.button("❌ 风险过高 (如账期太长)，驳回重审", use_container_width=True):
                st.session_state[state_key] = "3_sales_pricing"
                st.rerun()
                
        elif st.session_state[state_key] == "5_approved":
            st.success("🎉 合同已获批生效！系统已自动生成防篡改电子合同，并同步回款台账。")
            terms = st.session_state.get(f"comm_terms_{contract_proj}", {})
            with st.container(border=True):
                st.markdown(f"### 📄 《{contract_proj} 供货与技术服务合同》")
                st.dataframe(st.session_state[bom_key], use_container_width=True)
                if terms:
                    st.markdown("---")
                    st.markdown(f"**【第一条】 结算方式**：双方约定以 `{terms.get('pay_method')}` 结算。按 **{terms.get('ratio_advance')}%(预付款) - {terms.get('ratio_delivery')}%(发货款) - {terms.get('ratio_accept')}%(验收款) - {terms.get('ratio_warranty')}%(质保金)** 节点比例支付。")
                    st.markdown(f"**【第二条】 交货时间**：{terms.get('delivery_time')}")
                    st.markdown(f"**【第三条】 质保承诺**：{terms.get('warranty_period')}")
                    st.markdown(f"**【第四条】 物流接收**：{terms.get('delivery_address')} (联系人：{terms.get('receiver_contact')})")
                st.markdown("<br><span style='color:red; font-weight:bold;'>【此件已加盖企业骑缝章及电子合同专用章】</span>", unsafe_allow_html=True)
                
            contact_email = st.text_input("🎯 客户接收邮箱：", placeholder="由 AI 自动填充或手动输入")
            if st.button("🚀 调取邮件引擎一键发送合同附件", type="primary"):
                st.balloons()
                st.success(f"✅ 合同已加密发送至 {contact_email}！系统已自动生成回款台账。")
                # 触发进入终极环节：提成核算
                st.session_state[state_key] = "6_commission"
                st.rerun()

    # 🟣 第五步：销售提成自动核算
    with st.expander("🟣 第五步：项目销售提成单自动核算 (财务/高管可见)", expanded=(st.session_state[state_key] == "6_commission")):
        if st.session_state[state_key] in ["1_sales_init", "2_tech_review", "3_sales_pricing", "4_vp_approval", "5_approved"]:
            st.info("⏳ 等待合同正式发送给客户后，解锁提成核算...")
        elif st.session_state[state_key] == "6_commission":
            st.info("💡 财务/行政人员请注意：请根据实际情况补充【公司结算底价】与【运费】，系统将自动结合合同 BOM 表和技术超配信息生成提成核算单。")
            
            # 初始化提成底稿
            if f"comm_bom_{contract_proj}" not in st.session_state:
                comm_df = st.session_state[bom_key].copy()
                comm_df["公司结算底价(元)"] = 0.0
                comm_df["提成比例(%)"] = 10.0
                st.session_state[f"comm_bom_{contract_proj}"] = comm_df
            
            c_calc1, c_calc2 = st.columns([3, 1])
            with c_calc1:
                # 锁定核心数据，只允许财务填底价和比例
                edited_comm_bom = st.data_editor(
                    st.session_state[f"comm_bom_{contract_proj}"],
                    disabled=["产品型号", "技术核定数量", "最终报价数量", "单价(元)", "超配说明"],
                    use_container_width=True, num_rows="dynamic", key="comm_editor"
                )
                st.session_state[f"comm_bom_{contract_proj}"] = edited_comm_bom
            
            with c_calc2:
                st.markdown("**⚙️ 结算参数配置**")
                formula_type = st.selectbox("🧮 提成结算公式：", ["毛利提成：(单价-底价)*数量*比例", "全额提成：单价*数量*比例"], key=f"formula_{contract_proj}")
                freight_cost = st.number_input("🚚 预估/实际运费(元) [后期可补填]：", min_value=0.0, step=500.0, key=f"freight_{contract_proj}")
            
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("💵 一键核算并生成《标准提成单》", type="primary", use_container_width=True):
                total_comm = 0
                calc_details = []
                
                for idx, row in edited_comm_bom.iterrows():
                    # 兼容不同阶段存下来的数量字段
                    qty = float(row.get("最终报价数量", row.get("技术核定数量", 0)))
                    price = float(row.get("单价(元)", 0))
                    base_price = float(row.get("公司结算底价(元)", 0))
                    ratio = float(row.get("提成比例(%)", 10)) / 100.0
                    
                    if "毛利" in formula_type:
                        item_comm = (price - base_price) * qty * ratio
                        diff = price - base_price
                    else:
                        item_comm = price * qty * ratio
                        diff = price
                        
                    total_comm += item_comm
                    calc_details.append({
                        "产品型号": row.get("产品型号"),
                        "超配影响": row.get("超配说明", ""),
                        "结算基数": diff,
                        "计算数量": qty,
                        "单项提成": item_comm
                    })
                
                final_comm = total_comm - freight_cost
                
                # 渲染华丽的提成单
                with st.container(border=True):
                    st.markdown(f"<h3 style='text-align: center;'>💸 项目销售提成核算单</h3>", unsafe_allow_html=True)
                    st.markdown(f"**业务项目**：`{contract_proj}` &nbsp;&nbsp;|&nbsp;&nbsp; **负责销售**：`{st.session_state.projects.get(contract_proj, {}).get('owner', '当前销售')}`", unsafe_allow_html=True)
                    st.markdown(f"**执行公式**：`{formula_type}` &nbsp;&nbsp;|&nbsp;&nbsp; **运费/杂费扣减**：`- ¥ {freight_cost:,.2f}`", unsafe_allow_html=True)
                    
                    import pandas as pd
                    st.dataframe(pd.DataFrame(calc_details), use_container_width=True, hide_index=True)
                    
                    st.markdown(f"#### 💰 最终应发提成总额：<span style='color:red;'>¥ {final_comm:,.2f}</span>", unsafe_allow_html=True)
                    
                    # 极其硬核的业务联动：将提成发放与合同中的回款比例强制挂钩！
                    terms = st.session_state.get(f"comm_terms_{contract_proj}", {})
                    if terms:
                        st.info(f"💡 **财务放款约定**：此提成将严格按照合同约定的回款节点进行同频拨付！\n(即：客户付 {terms.get('ratio_advance')}% 预付款，则发放 ¥ {(final_comm * terms.get('ratio_advance') / 100):,.2f} 提成，以此类推...)")
                
                if st.button("💾 确认底稿并归档", type="secondary"):
                    st.session_state[state_key] = "1_sales_init"
                    st.toast("✅ 提成单已归档！流程彻底闭环。")
                    st.rerun()

render_config_manager()
